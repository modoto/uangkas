"""
Generate FSD (Functional Specification Document) for Uang Kas v1.0
Run: python generate_fsd.py
Output: FSD_UangKas_v1.0.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table_header(table, headers, bg='1E3A5F', fg='FFFFFF', font_size=9):
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(fg)

def add_table_row(table, values, bg=None, bold=False, align_center_cols=None, font_size=9):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        if align_center_cols and i in align_center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(font_size)
    return row

def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    if level == 1:
        para.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    elif level == 2:
        para.runs[0].font.color.rgb = RGBColor(0x2D, 0x5F, 0xA8)
    elif level >= 3:
        para.runs[0].font.color.rgb = RGBColor(0x1A, 0x7A, 0x5A)
    return para

def add_para(doc, text, bold=False, italic=False, size=10, space_after=4, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_numbered(doc, text, size=10):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2D5FA8')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_label(doc, text, bg='DCE6F1'):
    """Blok label section berwarna"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'  {text}  ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p

def field_table(doc, fields):
    """Tabel spesifikasi field form"""
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    add_table_header(t, ['Field', 'Tipe Input', 'Wajib', 'Validasi', 'Keterangan'], bg='2D5FA8')
    for f in fields:
        add_table_row(t, f, font_size=8.5)
    doc.add_paragraph()
    return t

# ─── Document Setup ──────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FUNCTIONAL SPECIFICATION DOCUMENT')
r.bold = True; r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UANG KAS')
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x2D, 0x5F, 0xA8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versi 1.0')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Aplikasi Monitoring Keuangan Operasional Driver')
r.italic = True; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_table(rows=6, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = 'Table Grid'
info_data = [
    ('Nama Dokumen',  'Functional Specification Document — Uang Kas'),
    ('Versi',         '1.0'),
    ('Tanggal',       '16 April 2026'),
    ('Status',        'Draft'),
    ('Dibuat Oleh',   'Tim Pengembang'),
    ('Ditujukan',     'Owner & Tim Pengembang'),
]
for i, (k, v) in enumerate(info_data):
    row = info.rows[i]
    set_cell_bg(row.cells[0], 'DCE6F1')
    row.cells[0].text = k
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = v
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# DAFTAR ISI
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Daftar Isi', 1)
hr(doc)

toc_items = [
    ('1.', 'Informasi Dokumen'),
    ('2.', 'Tujuan & Ruang Lingkup FSD'),
    ('3.', 'Arsitektur Sistem'),
    ('4.', 'Modul Autentikasi'),
    ('5.', 'Modul Owner — Dashboard'),
    ('6.', 'Modul Owner — Manajemen Driver'),
    ('7.', 'Modul Owner — Manajemen Wallet'),
    ('8.', 'Modul Owner — Jenis Transaksi'),
    ('9.', 'Modul Owner — Laporan & Monitoring'),
    ('10.', 'Modul Owner — Kelola Session'),
    ('11.', 'Modul Driver — Dashboard'),
    ('12.', 'Modul Driver — Transaksi Jenius'),
    ('13.', 'Modul Driver — Pengeluaran Cash'),
    ('14.', 'Modul Driver — Riwayat Transaksi'),
    ('15.', 'Spesifikasi Database Lengkap'),
    ('16.', 'Spesifikasi Email'),
    ('17.', 'Penanganan Error & Validasi'),
    ('18.', 'Spesifikasi Keamanan Teknis'),
    ('19.', 'Spesifikasi UI/UX'),
    ('20.', 'Skenario Pengujian Fungsional'),
]
toc_t = doc.add_table(rows=len(toc_items), cols=2)
toc_t.style = 'Table Grid'
for i, (num, title) in enumerate(toc_items):
    row = toc_t.rows[i]
    if i % 2 == 0:
        set_cell_bg(row.cells[0], 'F0F4FF')
        set_cell_bg(row.cells[1], 'F0F4FF')
    row.cells[0].text = num
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = title
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INFORMASI DOKUMEN
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Informasi Dokumen', 1)
hr(doc)

add_heading(doc, '1.1 Riwayat Revisi', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Versi', 'Tanggal', 'Perubahan', 'Oleh', 'Status'])
add_table_row(t, ['1.0', '16 Apr 2026', 'Dokumen awal — FSD Uang Kas', 'Tim Dev', 'Draft'], align_center_cols=[0,1,4])

doc.add_paragraph()
add_heading(doc, '1.2 Referensi Dokumen', 2)
add_bullet(doc, 'BRD Uang Kas v1.0 — Business Requirement Document (dokumen pendamping)')
add_bullet(doc, 'Flask Documentation — https://flask.palletsprojects.com/')
add_bullet(doc, 'SQLAlchemy Documentation — https://docs.sqlalchemy.org/')
add_bullet(doc, 'HTMX Documentation — https://htmx.org/docs/')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TUJUAN FSD
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Tujuan & Ruang Lingkup FSD', 1)
hr(doc)

add_para(doc, 'Dokumen ini menjabarkan spesifikasi fungsional teknis aplikasi Uang Kas v1.0 secara rinci. '
         'Dokumen ini digunakan oleh tim pengembang sebagai acuan implementasi dan oleh owner '
         'sebagai referensi verifikasi bahwa sistem yang dibangun sesuai dengan kebutuhan bisnis.')

add_para(doc, 'Perbedaan dengan BRD:', bold=True)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Aspek', 'BRD', 'FSD'])
diff_data = [
    ('Fokus',        'APA yang dibutuhkan bisnis',                    'BAGAIMANA sistem dibangun'),
    ('Pembaca',      'Owner, stakeholder bisnis, tim dev',             'Tim dev, QA, tech lead'),
    ('Level Detail', 'Kebutuhan bisnis & aturan bisnis',              'Spesifikasi layar, field, logika, validasi'),
    ('Output',       'Daftar requirement yang disetujui',             'Blueprint implementasi teknis'),
]
for row in diff_data:
    add_table_row(t, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ARSITEKTUR SISTEM
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Arsitektur Sistem', 1)
hr(doc)

add_heading(doc, '3.1 Arsitektur Aplikasi', 2)
add_para(doc, 'Uang Kas menggunakan arsitektur Monolithic MPA (Multi-Page Application) dengan pola '
         'request-response tradisional yang diperkuat HTMX untuk partial update tanpa full page reload.')

add_para(doc, 'Alur Request:', bold=True)
add_numbered(doc, 'Browser mengirim HTTP request ke Flask server.')
add_numbered(doc, 'Flask middleware memvalidasi session cookie (HTTPONLY + SAMESITE=Lax).')
add_numbered(doc, 'Route handler memvalidasi peran pengguna (owner/driver).')
add_numbered(doc, 'Business logic dieksekusi, database diakses via SQLAlchemy ORM.')
add_numbered(doc, 'Jinja2 me-render template HTML dengan data dari server.')
add_numbered(doc, 'Response HTML dikirim ke browser. Untuk request HTMX, hanya partial HTML yang dikirim.')

doc.add_paragraph()
add_heading(doc, '3.2 Stack Teknologi Detail', 2)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Layer', 'Teknologi', 'Versi', 'Fungsi Spesifik'])
tech_data = [
    ('Web Framework',   'Python Flask',             '3.x',     'Routing, request handling, session'),
    ('ORM',             'Flask-SQLAlchemy',         '3.x',     'Model, query, migration'),
    ('DB Driver',       'psycopg2-binary',          '2.9.x',   'Koneksi PostgreSQL'),
    ('DB Dev',          'SQLite (built-in Python)', '3.x',     'Development tanpa install'),
    ('DB Prod',         'PostgreSQL',               '15+',     'Production database'),
    ('Template Engine', 'Jinja2',                   '3.x',     'Auto-escape, template inheritance'),
    ('Partial Update',  'HTMX',                     '1.9.x',   'hx-post, hx-swap, hx-target'),
    ('Email',           'Flask-Mail',               '0.9.x',   'SMTP, pengiriman magic link'),
    ('Hashing',         'hashlib (built-in)',        'stdlib',  'SHA-256 untuk token magic link'),
    ('CSRF',            'Flask-WTF',                '1.x',     'CSRFProtect, csrf_token()'),
    ('Font',            'Google Fonts CDN',          '-',       'Plus Jakarta Sans, JetBrains Mono'),
    ('Frontend',        'HTML5/CSS3/JS Vanilla',    '-',        'Tanpa framework CSS/JS'),
]
for row in tech_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '3.3 Struktur Blueprint Flask', 2)
add_para(doc, 'Aplikasi menggunakan Flask Blueprint untuk memisahkan modul. Setiap Blueprint memiliki prefix URL dan template folder sendiri.')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Blueprint', 'Prefix URL', 'File', 'Modul'])
bp_data = [
    ('auth',            '/auth',            'routes/auth.py',              'Login, verify, logout'),
    ('owner_dashboard', '/owner',           'routes/owner/dashboard.py',   'Dashboard owner'),
    ('owner_drivers',   '/owner/drivers',   'routes/owner/drivers.py',     'CRUD driver'),
    ('owner_wallets',   '/owner/wallets',   'routes/owner/wallets.py',     'Topup, kelola wallet'),
    ('owner_txntypes',  '/owner/txn-types', 'routes/owner/txn_types.py',   'Jenis transaksi'),
    ('owner_reports',   '/owner/reports',   'routes/owner/reports.py',     'Laporan & monitoring'),
    ('owner_sessions',  '/owner/sessions',  'routes/owner/sessions.py',    'Kelola session driver'),
    ('driver_dashboard','/driver',          'routes/driver/dashboard.py',  'Dashboard driver'),
    ('driver_jenius',   '/driver/jenius',   'routes/driver/jenius.py',     'Transaksi Jenius'),
    ('driver_cash',     '/driver/cash',     'routes/driver/cash.py',       'Pengeluaran cash'),
    ('driver_history',  '/driver/history',  'routes/driver/history.py',    'Riwayat transaksi'),
]
for row in bp_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '3.4 Middleware & Decorator', 2)
add_para(doc, 'Dua decorator kustom digunakan untuk proteksi route:', )
add_bullet(doc, '@login_required — Memvalidasi session cookie ada dan aktif di tabel user_sessions. Jika tidak valid, redirect ke /auth/login.')
add_bullet(doc, '@owner_required — Memanggil @login_required terlebih dahulu, lalu memvalidasi current_user.role == "owner". Jika bukan owner, abort(403).')
add_bullet(doc, '@driver_required — Memanggil @login_required, lalu memvalidasi current_user.role == "driver". Jika bukan driver, abort(403).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MODUL AUTENTIKASI
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Modul Autentikasi', 1)
hr(doc)

# ── 4.1 Halaman Login ────────────────────────────────────────────────────────
add_heading(doc, '4.1 Halaman Login  [GET/POST /auth/login]', 2)

add_heading(doc, 'Deskripsi', 3)
add_para(doc, 'Halaman pertama yang dilihat pengguna. Pengguna memasukkan email untuk menerima magic link. '
         'Tidak ada input password — autentikasi sepenuhnya berbasis magic link.')

add_heading(doc, 'Elemen UI', 3)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Elemen', 'Tipe', 'Posisi', 'Keterangan'], bg='2D5FA8')
ui_data = [
    ('Logo "Uang Kas"',         'Teks/Gambar',   'Tengah atas',      'Nama aplikasi, ukuran besar'),
    ('Judul halaman',           'H1',            'Tengah',           '"Masuk ke Akun Anda"'),
    ('Subjudul',                'Paragraf',      'Bawah judul',      '"Kami akan mengirimkan link masuk ke email Anda"'),
    ('Input Email',             'input[email]',  'Tengah',           'Placeholder: "Alamat email Anda"'),
    ('Tombol "Kirim Magic Link"','button[submit]','Bawah input',     'Full width, warna primer'),
    ('CSRF Token',              'input[hidden]', 'Dalam form',       'Token proteksi CSRF'),
    ('Pesan error/sukses',      'div alert',     'Atas form',        'Muncul jika ada validasi gagal atau sukses'),
]
for row in ui_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, 'Logika Backend', 3)
add_numbered(doc, 'GET /auth/login → render template login.html. Jika user sudah punya session aktif, redirect ke dashboard sesuai peran.')
add_numbered(doc, 'POST /auth/login → ambil email dari form.')
add_numbered(doc, 'Query: users WHERE email = ? AND is_active = TRUE.')
add_numbered(doc, 'Jika tidak ditemukan: tampilkan pesan generik "Jika email terdaftar, link akan dikirimkan" (tidak membocorkan status).')
add_numbered(doc, 'Jika ditemukan: buat token random (secrets.token_urlsafe(32)), hash dengan SHA-256, simpan ke magic_tokens.')
add_numbered(doc, 'Rate limit: cek jumlah magic_token dengan user_id yang sama dalam 5 menit terakhir. Jika ≥ 3, tolak dan tampilkan pesan "Terlalu banyak permintaan, coba lagi nanti".')
add_numbered(doc, 'Kirim email berisi URL: https://[domain]/auth/verify/[token_plaintext].')
add_numbered(doc, 'Render halaman konfirmasi: "Link telah dikirimkan ke email Anda".')

doc.add_paragraph()
add_heading(doc, 'Validasi Field', 3)
field_table(doc, [
    ('email', 'input[email]', 'Ya', 'Format email valid; max 255 karakter', 'Input type=email + validasi server'),
])

# ── 4.2 Verifikasi Magic Link ─────────────────────────────────────────────────
add_heading(doc, '4.2 Verifikasi Magic Link  [GET /auth/verify/<token>]', 2)

add_heading(doc, 'Logika Backend', 3)
add_numbered(doc, 'Ambil token dari URL parameter.')
add_numbered(doc, 'Hash token dengan SHA-256.')
add_numbered(doc, 'Query: magic_tokens WHERE token = [hashed] AND used_at IS NULL AND expires_at > NOW().')
add_numbered(doc, 'Jika tidak ditemukan atau kedaluwarsa: redirect ke /auth/login dengan flash "Link tidak valid atau sudah kedaluwarsa".')
add_numbered(doc, 'Jika ditemukan: set magic_tokens.used_at = NOW().')
add_numbered(doc, 'Validasi user masih aktif: users WHERE id = token.user_id AND is_active = TRUE.')
add_numbered(doc, 'Jika user tidak aktif: redirect ke /auth/login dengan flash "Akun tidak aktif, hubungi administrator".')
add_numbered(doc, 'Buat session: INSERT user_sessions (user_id, session_token=secrets.token_urlsafe(64), device_info, ip_address).')
add_numbered(doc, 'Set Flask session cookie: session["session_token"] = [token], httponly=True, samesite="Lax", secure=[True jika production].')
add_numbered(doc, 'Redirect ke /owner/ jika role=owner, atau /driver/ jika role=driver.')

doc.add_paragraph()
add_heading(doc, '4.3 Logout  [POST /auth/logout]', 2)
add_heading(doc, 'Logika Backend', 3)
add_numbered(doc, 'Ambil session_token dari Flask session cookie.')
add_numbered(doc, 'UPDATE user_sessions SET is_active=FALSE, terminated_at=NOW() WHERE session_token = ?.')
add_numbered(doc, 'Hapus Flask session (session.clear()).')
add_numbered(doc, 'Redirect ke /auth/login dengan flash "Anda telah berhasil keluar".')

doc.add_paragraph()
add_heading(doc, '4.4 Mekanisme Validasi Session per Request', 2)
add_para(doc, 'Decorator @login_required dieksekusi sebelum setiap route yang membutuhkan autentikasi:')
add_numbered(doc, 'Ambil session["session_token"] dari cookie.')
add_numbered(doc, 'Jika tidak ada: redirect ke /auth/login.')
add_numbered(doc, 'Query: user_sessions JOIN users WHERE session_token = ? AND user_sessions.is_active = TRUE AND users.is_active = TRUE.')
add_numbered(doc, 'Jika tidak ditemukan (session diterminasi atau user dinonaktifkan): hapus cookie, redirect ke /auth/login dengan flash "Session tidak valid".')
add_numbered(doc, 'Set flask.g.current_user = user object untuk digunakan di route handler dan template.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OWNER DASHBOARD
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Modul Owner — Dashboard', 1)
hr(doc)

add_heading(doc, '5.1 Halaman Dashboard Owner  [GET /owner/]', 2)
add_heading(doc, 'Deskripsi', 3)
add_para(doc, 'Halaman utama owner setelah login. Memberikan ringkasan menyeluruh kondisi keuangan semua driver.')

add_heading(doc, 'Komponen Halaman', 3)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Komponen', 'Data yang Ditampilkan', 'Sumber Data'], bg='2D5FA8')
comp_data = [
    ('Header Navigasi',         'Logo, nama owner, menu navigasi, tombol logout',
     'g.current_user'),
    ('Kartu Ringkasan per Driver','Nama driver + 3 saldo (Jenius, eMoney, Cash) dalam kartu terpisah',
     'wallets JOIN users WHERE role=driver'),
    ('Total Semua Wallet',       'Jumlah total semua saldo (Jenius + eMoney + Cash) seluruh driver',
     'SUM(wallets.current_balance)'),
    ('Tabel 10 Transaksi Terakhir','Tanggal, Driver, Wallet, Jenis, Nominal, Saldo Akhir',
     'transactions JOIN wallets JOIN users ORDER BY created_at DESC LIMIT 10'),
    ('Indikator Status Driver',  'Badge "Aktif"/"Nonaktif" per driver',
     'users.is_active'),
]
for row in comp_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, 'Query Utama', 3)
add_code(doc, '# Ambil semua driver dengan wallet mereka')
add_code(doc, 'drivers = User.query.filter_by(role="driver").order_by(User.name).all()')
add_code(doc, '')
add_code(doc, '# Untuk setiap driver, ambil saldo wallet')
add_code(doc, 'for driver in drivers:')
add_code(doc, '    driver.wallets = Wallet.query.filter_by(driver_id=driver.id, is_active=True).all()')
add_code(doc, '')
add_code(doc, '# 10 transaksi terakhir dari semua driver')
add_code(doc, 'recent_txns = db.session.query(Transaction, Wallet, User)')
add_code(doc, '    .join(Wallet, Transaction.wallet_id == Wallet.id)')
add_code(doc, '    .join(User, Transaction.created_by == User.id)')
add_code(doc, '    .order_by(Transaction.created_at.desc()).limit(10).all()')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — MANAJEMEN DRIVER
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Modul Owner — Manajemen Driver', 1)
hr(doc)

add_heading(doc, '6.1 Daftar Driver  [GET /owner/drivers]', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Kolom Tabel', 'Sumber Data', 'Aksi'], bg='2D5FA8')
drv_list = [
    ('No.',             'Index',                          '-'),
    ('Nama Driver',     'users.name',                     'Link ke detail'),
    ('Email',           'users.email',                    '-'),
    ('Status',          'users.is_active',                'Badge Aktif/Nonaktif'),
    ('Saldo Jenius',    'wallets WHERE type=jenius',       '-'),
    ('Saldo eMoney',    'wallets WHERE type=emoney',       '-'),
    ('Saldo Cash',      'wallets WHERE type=cash',         '-'),
    ('Terdaftar',       'users.created_at',               '-'),
    ('Aksi',            '-',                              'Tombol Edit, Nonaktifkan/Aktifkan'),
]
for row in drv_list:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '6.2 Tambah Driver  [GET/POST /owner/drivers/add]', 2)
add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Nama Lengkap', 'text',  'Ya', 'Min 2 karakter, max 100 karakter',    'Nama yang ditampilkan di aplikasi'),
    ('Email',        'email', 'Ya', 'Format email valid, belum terdaftar',  'Email untuk magic link onboarding'),
])
add_heading(doc, 'Logika Backend POST', 3)
add_numbered(doc, 'Validasi: email belum terdaftar di tabel users.')
add_numbered(doc, 'INSERT ke users: name, email, role="driver", is_active=TRUE.')
add_numbered(doc, 'INSERT 3 wallet otomatis: Jenius (wallet_type="jenius"), eMoney (wallet_type="emoney"), Cash (wallet_type="cash"), semua current_balance=0.')
add_numbered(doc, 'Buat token onboarding: secrets.token_urlsafe(32), hash SHA-256, INSERT ke magic_tokens (purpose="onboarding", expires_at=NOW()+24jam).')
add_numbered(doc, 'Kirim email onboarding ke driver (lihat Seksi 16 — Spesifikasi Email).')
add_numbered(doc, 'Redirect ke /owner/drivers dengan flash "Driver berhasil ditambahkan, email onboarding telah dikirim".')

doc.add_paragraph()
add_heading(doc, '6.3 Edit Driver  [GET/POST /owner/drivers/<id>/edit]', 2)
add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Nama Lengkap',  'text',     'Ya',  'Min 2, max 100 karakter',     'Dapat diubah'),
    ('Status Aktif',  'checkbox', 'Ya',  '-',                           'Centang = aktif, kosong = nonaktif'),
])
add_heading(doc, 'Logika Backend POST', 3)
add_numbered(doc, 'UPDATE users SET name=?, is_active=? WHERE id=?.')
add_numbered(doc, 'Jika is_active berubah dari TRUE ke FALSE: UPDATE user_sessions SET is_active=FALSE, terminated_at=NOW(), terminated_by=owner_id WHERE user_id=driver_id AND is_active=TRUE.')
add_numbered(doc, 'Redirect ke /owner/drivers dengan flash sukses.')

doc.add_paragraph()
add_heading(doc, '6.4 Kirim Ulang Magic Link  [POST /owner/drivers/<id>/resend]', 2)
add_numbered(doc, 'Validasi: driver dengan id tersebut ada dan is_active=TRUE.')
add_numbered(doc, 'Invalidasi token onboarding sebelumnya: UPDATE magic_tokens SET used_at=NOW() WHERE user_id=? AND purpose="onboarding" AND used_at IS NULL.')
add_numbered(doc, 'Buat token onboarding baru dan kirim email.')
add_numbered(doc, 'Redirect dengan flash sukses.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — MANAJEMEN WALLET
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Modul Owner — Manajemen Wallet', 1)
hr(doc)

add_heading(doc, '7.1 Daftar Wallet  [GET /owner/wallets]', 2)
add_para(doc, 'Menampilkan semua wallet dikelompokkan per driver. Setiap grup menampilkan nama driver dan tiga kartu wallet (Jenius, eMoney, Cash) dengan saldo terkini.')

doc.add_paragraph()
add_heading(doc, '7.2 Top Up Jenius  [POST /owner/wallets/<id>/topup]', 2)
add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Driver',           'select',   'Ya',  'Pilih dari daftar driver aktif',              'Dropdown, pre-select jika dari konteks driver'),
    ('Nominal (Rp)',     'number',   'Ya',  'Integer positif, min 1000, max 999.999.999',  'Input angka tanpa desimal'),
    ('Tanggal & Jam',    'datetime', 'Ya',  'Tidak boleh di masa depan',                   'Default: waktu sekarang'),
    ('Catatan',          'textarea', 'Tidak','Max 500 karakter',                           'Keterangan topup opsional'),
])
add_heading(doc, 'Logika Backend', 3)
add_numbered(doc, 'Validasi wallet: wallet.id ada, wallet.wallet_type == "jenius", wallet.driver_id sesuai driver yang dipilih.')
add_numbered(doc, 'Mulai database transaction.')
add_numbered(doc, 'Hitung balance_before = wallet.current_balance.')
add_numbered(doc, 'Hitung balance_after = balance_before + nominal.')
add_numbered(doc, 'INSERT transactions: wallet_id, txn_type_id=[id dari code="owner_topup"], created_by=owner_id, direction="credit", amount=nominal, balance_before, balance_after, notes, txn_date.')
add_numbered(doc, 'UPDATE wallets SET current_balance = balance_after WHERE id = wallet_id.')
add_numbered(doc, 'COMMIT.')
add_numbered(doc, 'Redirect ke /owner/wallets dengan flash sukses.')

doc.add_paragraph()
add_heading(doc, '7.3 Top Up eMoney  [POST /owner/wallets/<id>/topup-emoney]', 2)
add_para(doc, 'Alur identik dengan Top Up Jenius (7.2). Perbedaan: txn_type.code = "owner_topup" untuk wallet_type="emoney".')

doc.add_paragraph()
add_heading(doc, '7.4 Penarikan Jenius (Owner)  [POST /owner/wallets/<id>/withdraw]', 2)
add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Nominal (Rp)',     'number',   'Ya',  'Integer positif, tidak melebihi saldo tersedia',  'Validasi saldo server-side'),
    ('Tanggal & Jam',    'datetime', 'Ya',  'Tidak boleh di masa depan',                        'Default: sekarang'),
    ('Catatan',          'textarea', 'Tidak','Max 500 karakter',                                'Alasan penarikan'),
])
add_heading(doc, 'Validasi Tambahan', 3)
add_bullet(doc, 'Cek: wallet.current_balance >= nominal. Jika tidak, kembalikan error "Saldo tidak mencukupi" dengan saldo terkini.')
add_bullet(doc, 'Transaksi: direction="debit", txn_type.code="owner_withdrawal".')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — JENIS TRANSAKSI
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '8. Modul Owner — Jenis Transaksi', 1)
hr(doc)

add_heading(doc, '8.1 Daftar Jenis Transaksi  [GET /owner/txn-types]', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Kolom Tabel', 'Sumber Data', 'Keterangan'], bg='2D5FA8')
tt_cols_disp = [
    ('Kode',         'transaction_types.code',        'Tidak dapat diubah untuk is_system=TRUE'),
    ('Label',        'transaction_types.label',        'Nama tampilan di dropdown driver'),
    ('Berlaku untuk','transaction_types.wallet_type',  'jenius / emoney / cash / all'),
    ('Arah',         'transaction_types.direction',    'credit / debit / both'),
    ('Urutan',       'transaction_types.sort_order',   'Angka kecil = tampil lebih awal'),
    ('System',       'transaction_types.is_system',    'Badge "System" jika TRUE'),
    ('Status',       'transaction_types.is_active',    'Toggle aktif/nonaktif'),
    ('Aksi',         '-',                              'Edit; Hapus (hanya non-system yang belum dipakai)'),
]
for row in tt_cols_disp:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '8.2 Tambah Jenis Transaksi  [GET/POST /owner/txn-types/add]', 2)
field_table(doc, [
    ('Label Tampilan',  'text',    'Ya',  'Min 3, max 100 karakter; tidak duplikat',         'e.g. "Bayar Cuci Mobil"'),
    ('Berlaku untuk',   'select',  'Ya',  'Pilih: jenius / emoney / cash / all',             'Menentukan di wallet mana jenis ini muncul'),
    ('Arah Default',    'select',  'Ya',  'Pilih: credit / debit / both',                   'Menentukan efek ke saldo'),
    ('Urutan Tampil',   'number',  'Tidak','Integer 1–999; default 99',                     'Urutan di dropdown driver'),
])
add_heading(doc, 'Logika Backend', 3)
add_numbered(doc, 'Validasi: label belum dipakai di transaction_types.')
add_numbered(doc, 'Generate code dari label: lower(), replace(" ", "_"), hapus karakter non-alfanumerik.')
add_numbered(doc, 'Jika code sudah ada: tambahkan suffix angka (_2, _3, dst).')
add_numbered(doc, 'INSERT transaction_types: code, label, wallet_type, direction, is_system=FALSE, is_active=TRUE, sort_order, created_by=owner_id.')

doc.add_paragraph()
add_heading(doc, '8.3 Edit Jenis Transaksi  [GET/POST /owner/txn-types/<id>/edit]', 2)
add_para(doc, 'Field yang dapat diedit berbeda berdasarkan is_system:')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Field', 'is_system = TRUE', 'is_system = FALSE'], bg='2D5FA8')
edit_rules = [
    ('Label',        'Dapat diedit',  'Dapat diedit'),
    ('Wallet Type',  'READ ONLY',     'Dapat diedit'),
    ('Direction',    'READ ONLY',     'Dapat diedit'),
    ('Sort Order',   'Dapat diedit',  'Dapat diedit'),
    ('Code',         'READ ONLY',     'READ ONLY (auto-generated)'),
    ('is_active',    'Dapat diubah',  'Dapat diubah'),
]
for row in edit_rules:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '8.4 Hapus Jenis Transaksi  [POST /owner/txn-types/<id>/delete]', 2)
add_numbered(doc, 'Validasi is_system == FALSE. Jika TRUE: abort(403).')
add_numbered(doc, 'Validasi belum digunakan: SELECT COUNT(*) FROM transactions WHERE txn_type_id = id.')
add_numbered(doc, 'Jika sudah digunakan: return error "Tidak dapat menghapus, jenis transaksi sudah digunakan di [X] transaksi".')
add_numbered(doc, 'Jika belum digunakan: DELETE FROM transaction_types WHERE id = ?.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — LAPORAN & MONITORING
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '9. Modul Owner — Laporan & Monitoring', 1)
hr(doc)

add_heading(doc, '9.1 Laporan Transaksi Jenius  [GET /owner/reports/jenius]', 2)
add_heading(doc, 'Parameter Filter (Query String)', 3)
field_table(doc, [
    ('driver_id',    'select',  'Tidak', 'UUID valid atau "all"',          'Default: semua driver'),
    ('date_from',    'date',    'Tidak', 'Format YYYY-MM-DD',              'Default: awal bulan berjalan'),
    ('date_to',      'date',    'Tidak', 'Format YYYY-MM-DD, >= date_from','Default: hari ini'),
    ('txn_type',     'select',  'Tidak', 'ID transaction_type atau "all"', 'Default: semua jenis'),
])
add_heading(doc, 'Kolom Tabel Hasil', 3)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Sumber', 'Format'], bg='2D5FA8')
report_cols = [
    ('Tanggal & Jam',       'transactions.txn_date',        'DD/MM/YYYY HH:MM'),
    ('Driver',              'users.name',                   'Nama lengkap'),
    ('Jenis Transaksi',     'transaction_types.label',      'Label tampilan'),
    ('Nominal',             'transactions.amount',          'Rp xxx.xxx (format Rupiah)'),
    ('Kredit/Debit',        'transactions.direction',       'Badge hijau (credit) / merah (debit)'),
    ('Saldo Sebelum',       'transactions.balance_before',  'Rp xxx.xxx'),
    ('Saldo Sesudah',       'transactions.balance_after',   'Rp xxx.xxx'),
    ('Catatan',             'transactions.notes',           'Teks, kosong jika NULL'),
    ('Diinput oleh',        'users.name (created_by)',      'Nama owner/driver yang input'),
]
for row in report_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '9.2 Laporan Uang Cash  [GET /owner/reports/cash]', 2)
add_para(doc, 'Laporan ini menampilkan ringkasan cash per driver dalam format "rekening koran sederhana".')
add_heading(doc, 'Struktur Tampilan', 3)
add_numbered(doc, 'Header: nama driver, periode yang dipilih.')
add_numbered(doc, 'Saldo Awal Periode: saldo cash pada tanggal date_from (balance_after dari transaksi terakhir sebelum date_from, atau 0 jika belum ada transaksi).')
add_numbered(doc, 'Tabel transaksi cash: tanggal, keterangan (jenis + catatan), masuk (credit), keluar (debit), saldo berjalan.')
add_numbered(doc, 'Saldo Akhir Periode: saldo cash pada akhir date_to.')
add_numbered(doc, 'Ringkasan: total pemasukan (dari ATM), total pengeluaran, selisih.')

doc.add_paragraph()
add_heading(doc, '9.3 Rekap eMoney  [GET /owner/reports/emoney]', 2)
add_para(doc, 'Menampilkan riwayat top up eMoney oleh owner dan transaksi pemakaian yang dicatat driver (emoney_usage). '
         'Struktur tabel sama dengan laporan Jenius pada 9.1.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — KELOLA SESSION
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '10. Modul Owner — Kelola Session', 1)
hr(doc)

add_heading(doc, '10.1 Daftar Session Aktif  [GET /owner/sessions]', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Sumber Data', 'Keterangan'], bg='2D5FA8')
sess_cols = [
    ('Driver',          'users.name',                       'Nama driver pemilik session'),
    ('Email',           'users.email',                      'Email driver'),
    ('Perangkat',       'user_sessions.device_info',        'User-Agent browser/app'),
    ('IP Address',      'user_sessions.ip_address',         'IP saat login'),
    ('Login Sejak',     'user_sessions.created_at',         'Format: DD/MM/YYYY HH:MM'),
    ('Status',          'user_sessions.is_active',          'Hanya tampilkan is_active=TRUE'),
    ('Aksi',            '-',                                'Tombol "Terminate Session"'),
]
for row in sess_cols:
    add_table_row(t, list(row), font_size=9)
add_para(doc, 'Catatan: session milik owner yang sedang login tidak ditampilkan di daftar ini.', italic=True)

doc.add_paragraph()
add_heading(doc, '10.2 Terminate Session  [POST /owner/sessions/<id>/terminate]', 2)
add_numbered(doc, 'Validasi: session dengan id ada dan is_active=TRUE.')
add_numbered(doc, 'Validasi: session bukan milik owner (user_sessions.user_id != g.current_user.id).')
add_numbered(doc, 'UPDATE user_sessions SET is_active=FALSE, terminated_at=NOW(), terminated_by=owner_id WHERE id=?.')
add_numbered(doc, 'Redirect ke /owner/sessions dengan flash "Session driver [nama] berhasil diterminasi".')
add_numbered(doc, 'Pada request berikutnya dari driver tersebut: decorator @login_required mendeteksi is_active=FALSE, hapus cookie, redirect ke /auth/login.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DRIVER DASHBOARD
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '11. Modul Driver — Dashboard', 1)
hr(doc)

add_heading(doc, '11.1 Halaman Dashboard Driver  [GET /driver/]', 2)
add_heading(doc, 'Deskripsi', 3)
add_para(doc, 'Halaman utama driver, didesain mobile-first. Menampilkan saldo tiga wallet dan akses cepat ke fitur pencatatan transaksi.')

add_heading(doc, 'Komponen Halaman (Mobile Layout)', 3)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Komponen', 'Data', 'Keterangan'], bg='2D5FA8')
drv_dash = [
    ('Header',                   'Nama driver, tombol logout',           'Sticky header'),
    ('Kartu Jenius',             'Label wallet, saldo terkini',          'Warna biru, ikon kartu'),
    ('Kartu eMoney',             'Label wallet, saldo terkini',          'Warna hijau, ikon kartu e-money'),
    ('Kartu Uang Cash',          'Label wallet, saldo terkini',          'Warna oranye, ikon uang'),
    ('Tombol Aksi Cepat',        '"Catat Jenius" dan "Catat Cash"',      'Tombol besar, mudah diklik di mobile'),
    ('5 Transaksi Terakhir',     'Tanggal, jenis, nominal (merah/hijau)','Ringkasan aktivitas terbaru'),
]
for row in drv_dash:
    add_table_row(t, list(row), font_size=9)

add_heading(doc, 'Query Utama', 3)
add_code(doc, '# Wallet milik driver yang login')
add_code(doc, 'wallets = Wallet.query.filter_by(driver_id=current_user.id, is_active=True).all()')
add_code(doc, '')
add_code(doc, '# 5 transaksi terakhir driver')
add_code(doc, 'recent = Transaction.query')
add_code(doc, '    .join(Wallet).filter(Wallet.driver_id == current_user.id)')
add_code(doc, '    .order_by(Transaction.txn_date.desc()).limit(5).all()')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — TRANSAKSI JENIUS DRIVER
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '12. Modul Driver — Transaksi Jenius', 1)
hr(doc)

add_heading(doc, '12.1 Halaman Catat Transaksi Jenius  [GET/POST /driver/jenius/record]', 2)
add_heading(doc, 'Deskripsi', 3)
add_para(doc, 'Satu halaman untuk mencatat semua jenis transaksi pada wallet Jenius. Driver memilih jenis transaksi '
         'dari pilihan yang tersedia. Untuk Tarik ATM, sistem otomatis membuat dua transaksi terhubung (linked transaction).')

add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Jenis Transaksi',  'radio + select', 'Ya',
     'Pilihan Tarik ATM, Bayar Bensin, atau Lainnya (dropdown)',
     'Pilihan "Lainnya" menampilkan dropdown transaction_types aktif untuk wallet Jenius'),
    ('Nominal (Rp)',     'number',         'Ya',
     'Integer positif, min 1000, maks saldo Jenius tersedia (untuk debit)',
     'Validasi saldo dilakukan server-side'),
    ('Tanggal',          'date',           'Ya',
     'Format YYYY-MM-DD, tidak boleh di masa depan',
     'Default: hari ini'),
    ('Jam',              'time',           'Ya',
     'Format HH:MM',
     'Default: jam sekarang'),
    ('Catatan',          'textarea',       'Tidak',
     'Max 500 karakter',
     'Contoh: "SPBU Gatot Subroto", "tip parkir Sudirman"'),
])

add_heading(doc, 'Logika Backend — Pilihan "Tarik ATM" (Linked Transaction)', 3)
add_numbered(doc, 'Validasi: saldo Jenius driver >= nominal.')
add_numbered(doc, 'Ambil txn_type untuk code="atm_withdrawal" dan code="cash_in".')
add_numbered(doc, 'Ambil wallet Jenius dan wallet Cash milik driver yang login.')
add_numbered(doc, 'Mulai db.session.begin() (database transaction).')
add_numbered(doc, 'INSERT transactions (JENIUS): wallet_id=jenius.id, txn_type_id=atm_withdrawal.id, direction="debit", amount=nominal, balance_before=jenius.current_balance, balance_after=jenius.current_balance-nominal, txn_date, notes, created_by=driver.id.')
add_numbered(doc, 'UPDATE wallets SET current_balance = current_balance - nominal WHERE id = jenius.id.')
add_numbered(doc, 'INSERT transactions (CASH): wallet_id=cash.id, txn_type_id=cash_in.id, direction="credit", amount=nominal, balance_before=cash.current_balance, balance_after=cash.current_balance+nominal, txn_date=txn_date, notes="Dari Tarik ATM", created_by=driver.id, linked_txn_id=id_jenius_txn.')
add_numbered(doc, 'UPDATE transactions SET linked_txn_id = cash_txn.id WHERE id = jenius_txn.id. (update balik untuk dua arah)')
add_numbered(doc, 'UPDATE wallets SET current_balance = current_balance + nominal WHERE id = cash.id.')
add_numbered(doc, 'db.session.commit().')
add_numbered(doc, 'Redirect ke /driver/ dengan flash "Transaksi berhasil dicatat. Saldo Cash bertambah Rp [nominal]".')

add_heading(doc, 'Logika Backend — Pilihan "Bayar Bensin" atau "Lainnya"', 3)
add_numbered(doc, 'Validasi: saldo Jenius >= nominal (jika direction=debit).')
add_numbered(doc, 'INSERT satu transaksi saja ke wallet Jenius.')
add_numbered(doc, 'UPDATE saldo wallet Jenius.')
add_numbered(doc, 'COMMIT.')

add_heading(doc, 'Perilaku HTMX', 3)
add_bullet(doc, 'Ketika driver memilih radio "Lainnya", HTMX mengirim GET request ke /driver/jenius/txn-types untuk mengambil HTML dropdown jenis transaksi yang berlaku untuk wallet Jenius.')
add_bullet(doc, 'Dropdown di-inject ke dalam form tanpa full page reload menggunakan hx-swap="outerHTML".')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — PENGELUARAN CASH DRIVER
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '13. Modul Driver — Pengeluaran Cash', 1)
hr(doc)

add_heading(doc, '13.1 Halaman Catat Pengeluaran Cash  [GET/POST /driver/cash/record]', 2)
add_heading(doc, 'Deskripsi', 3)
add_para(doc, 'Driver mencatat setiap pengeluaran uang tunai. Catatan wajib diisi untuk transparansi '
         'penggunaan dana. Saldo cash dikurangi sesuai nominal.')

add_heading(doc, 'Spesifikasi Field Form', 3)
field_table(doc, [
    ('Nominal (Rp)',  'number',   'Ya',   'Integer positif, min 500, tidak melebihi saldo Cash',
     'Validasi saldo dilakukan server-side'),
    ('Tanggal',       'date',     'Ya',   'Tidak boleh di masa depan',  'Default: hari ini'),
    ('Jam',           'time',     'Ya',   'Format HH:MM',               'Default: jam sekarang'),
    ('Catatan',       'textarea', 'Ya',   'Min 3 karakter, max 500',    'Wajib: untuk apa uang digunakan'),
])

add_heading(doc, 'Logika Backend POST', 3)
add_numbered(doc, 'Ambil wallet Cash milik driver yang login.')
add_numbered(doc, 'Validasi saldo: wallet.current_balance >= nominal.')
add_numbered(doc, 'Ambil txn_type untuk code="cash_out".')
add_numbered(doc, 'Mulai database transaction.')
add_numbered(doc, 'INSERT transactions: wallet_id=cash.id, txn_type_id=cash_out.id, direction="debit", amount=nominal, balance_before, balance_after=balance_before-nominal, txn_date=tanggal+jam, notes=catatan, created_by=driver.id.')
add_numbered(doc, 'UPDATE wallets SET current_balance = current_balance - nominal WHERE id = cash.id.')
add_numbered(doc, 'COMMIT.')
add_numbered(doc, 'Redirect ke /driver/ dengan flash "Pengeluaran cash berhasil dicatat".')

add_heading(doc, 'Tampilan Saldo Tersedia', 3)
add_para(doc, 'Di atas form, tampilkan saldo cash saat ini secara dinamis:')
add_bullet(doc, 'Teks: "Saldo Cash Tersedia: Rp [saldo]" — diambil dari wallets.current_balance saat halaman dimuat.')
add_bullet(doc, 'Jika saldo = 0: tampilkan warning "Saldo cash kosong. Lakukan Tarik ATM terlebih dahulu."')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — RIWAYAT TRANSAKSI DRIVER
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '14. Modul Driver — Riwayat Transaksi', 1)
hr(doc)

add_heading(doc, '14.1 Riwayat Transaksi Jenius  [GET /driver/history/jenius]', 2)
add_para(doc, 'Menampilkan semua transaksi wallet Jenius milik driver yang login, urut dari terbaru.')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Keterangan'], bg='2D5FA8')
hist_data = [
    ('Tanggal & Jam',   'transactions.txn_date — format DD MMM YYYY, HH:MM'),
    ('Jenis',           'transaction_types.label'),
    ('Nominal',         'Rp [amount] — warna merah jika debit, hijau jika credit'),
    ('Saldo Setelah',   'transactions.balance_after — format Rp xxx.xxx'),
    ('Catatan',         'transactions.notes — disingkat jika panjang, tooltip untuk full text'),
    ('Linked',          'Ikon link jika linked_txn_id tidak NULL, tooltip "Terhubung ke transaksi Cash"'),
]
for row in hist_data:
    add_table_row(t, list(row), font_size=9)
add_para(doc, 'Pagination: 20 transaksi per halaman. Filter tanggal opsional.', italic=True)

doc.add_paragraph()
add_heading(doc, '14.2 Riwayat Transaksi Cash  [GET /driver/history/cash]', 2)
add_para(doc, 'Identik dengan 14.1 namun untuk wallet Cash. Penambahan kolom khusus:')
add_bullet(doc, 'Kolom "Sumber" pada baris cash_in: menampilkan "Dari Tarik ATM" dengan link ke transaksi Jenius terkait.')
add_bullet(doc, 'Ringkasan di atas tabel: Total Masuk (cash_in), Total Keluar (cash_out), Saldo Sekarang.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — SPESIFIKASI DATABASE LENGKAP
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '15. Spesifikasi Database Lengkap', 1)
hr(doc)

add_heading(doc, '15.1 Diagram Relasi (Tekstual)', 2)
erd_code = (
    "users\n"
    "  │\n"
    "  ├──(role=driver)── wallets ──────── transactions\n"
    "  │                     │                  │\n"
    "  ├── magic_tokens       │           transaction_types\n"
    "  │                      │                  │\n"
    "  ├── user_sessions       └── transactions ──┘\n"
    "  │                              │\n"
    "  └──(created_by)────────────────┘\n"
    "                   (self-ref: linked_txn_id)\n"
    "  transactions ◄────────────────── transactions\n"
    "  (atm_withdrawal)              (cash_in)"
)
add_code(doc, erd_code)

doc.add_paragraph()
add_heading(doc, '15.2 DDL PostgreSQL', 2)

ddl_blocks = [
    ('-- Tabel users', [
        "CREATE TABLE users (",
        "    id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    email       VARCHAR(255) UNIQUE NOT NULL,",
        "    name        VARCHAR(100) NOT NULL,",
        "    role        VARCHAR(10) NOT NULL CHECK (role IN ('owner','driver')),",
        "    is_active   BOOLEAN NOT NULL DEFAULT TRUE,",
        "    created_at  TIMESTAMP NOT NULL DEFAULT NOW(),",
        "    updated_at  TIMESTAMP NOT NULL DEFAULT NOW()",
        ");",
        "CREATE INDEX idx_users_email ON users(email);",
        "CREATE INDEX idx_users_role  ON users(role);",
    ]),
    ('-- Tabel magic_tokens', [
        "CREATE TABLE magic_tokens (",
        "    id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    user_id     UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,",
        "    token       VARCHAR(128) UNIQUE NOT NULL,",
        "    purpose     VARCHAR(20) NOT NULL CHECK (purpose IN ('login','onboarding')),",
        "    expires_at  TIMESTAMP NOT NULL,",
        "    used_at     TIMESTAMP,",
        "    created_at  TIMESTAMP NOT NULL DEFAULT NOW()",
        ");",
        "CREATE INDEX idx_magic_tokens_user ON magic_tokens(user_id);",
        "CREATE INDEX idx_magic_tokens_token ON magic_tokens(token);",
    ]),
    ('-- Tabel user_sessions', [
        "CREATE TABLE user_sessions (",
        "    id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    user_id         UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,",
        "    session_token   VARCHAR(255) UNIQUE NOT NULL,",
        "    device_info     VARCHAR(255),",
        "    ip_address      VARCHAR(45),",
        "    is_active       BOOLEAN NOT NULL DEFAULT TRUE,",
        "    created_at      TIMESTAMP NOT NULL DEFAULT NOW(),",
        "    terminated_at   TIMESTAMP,",
        "    terminated_by   UUID REFERENCES users(id)",
        ");",
        "CREATE INDEX idx_sessions_user      ON user_sessions(user_id);",
        "CREATE INDEX idx_sessions_token     ON user_sessions(session_token);",
        "CREATE INDEX idx_sessions_is_active ON user_sessions(is_active);",
    ]),
    ('-- Tabel transaction_types', [
        "CREATE TABLE transaction_types (",
        "    id           UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    code         VARCHAR(50) UNIQUE NOT NULL,",
        "    label        VARCHAR(100) NOT NULL,",
        "    wallet_type  VARCHAR(20) NOT NULL CHECK (wallet_type IN ('jenius','emoney','cash','all')),",
        "    direction    VARCHAR(10) NOT NULL CHECK (direction IN ('credit','debit','both')),",
        "    is_system    BOOLEAN NOT NULL DEFAULT FALSE,",
        "    is_active    BOOLEAN NOT NULL DEFAULT TRUE,",
        "    sort_order   SMALLINT NOT NULL DEFAULT 99,",
        "    created_by   UUID REFERENCES users(id),",
        "    created_at   TIMESTAMP NOT NULL DEFAULT NOW()",
        ");",
    ]),
    ('-- Tabel wallets', [
        "CREATE TABLE wallets (",
        "    id               UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    driver_id        UUID NOT NULL REFERENCES users(id) ON DELETE RESTRICT,",
        "    wallet_type      VARCHAR(20) NOT NULL CHECK (wallet_type IN ('jenius','emoney','cash')),",
        "    label            VARCHAR(100) NOT NULL,",
        "    current_balance  BIGINT NOT NULL DEFAULT 0 CHECK (current_balance >= 0),",
        "    is_active        BOOLEAN NOT NULL DEFAULT TRUE,",
        "    created_at       TIMESTAMP NOT NULL DEFAULT NOW(),",
        "    UNIQUE (driver_id, wallet_type)",
        ");",
        "CREATE INDEX idx_wallets_driver ON wallets(driver_id);",
    ]),
    ('-- Tabel transactions', [
        "CREATE TABLE transactions (",
        "    id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),",
        "    wallet_id       UUID NOT NULL REFERENCES wallets(id) ON DELETE RESTRICT,",
        "    txn_type_id     UUID NOT NULL REFERENCES transaction_types(id) ON DELETE RESTRICT,",
        "    created_by      UUID NOT NULL REFERENCES users(id) ON DELETE RESTRICT,",
        "    direction       VARCHAR(10) NOT NULL CHECK (direction IN ('credit','debit')),",
        "    amount          BIGINT NOT NULL CHECK (amount > 0),",
        "    balance_before  BIGINT NOT NULL,",
        "    balance_after   BIGINT NOT NULL,",
        "    notes           TEXT,",
        "    txn_date        TIMESTAMP NOT NULL,",
        "    linked_txn_id   UUID REFERENCES transactions(id),",
        "    created_at      TIMESTAMP NOT NULL DEFAULT NOW()",
        ");",
        "CREATE INDEX idx_txn_wallet     ON transactions(wallet_id);",
        "CREATE INDEX idx_txn_type       ON transactions(txn_type_id);",
        "CREATE INDEX idx_txn_created_by ON transactions(created_by);",
        "CREATE INDEX idx_txn_txn_date   ON transactions(txn_date);",
        "CREATE INDEX idx_txn_linked     ON transactions(linked_txn_id);",
    ]),
]

for comment, lines in ddl_blocks:
    add_para(doc, comment, bold=True, size=9, color='1A7A5A')
    for line in lines:
        add_code(doc, line)
    doc.add_paragraph()

add_heading(doc, '15.3 Data Seed transaction_types', 2)
seed_t = doc.add_table(rows=1, cols=5)
seed_t.style = 'Table Grid'
add_table_header(seed_t, ['code', 'label', 'wallet_type', 'direction', 'is_system'])
seed_rows = [
    ('owner_topup',      'Top Up Saldo',       'jenius',  'credit', 'TRUE'),
    ('owner_topup_em',   'Top Up eMoney',       'emoney',  'credit', 'TRUE'),
    ('owner_withdrawal', 'Penarikan Saldo',     'jenius',  'debit',  'TRUE'),
    ('atm_withdrawal',   'Tarik ATM',           'jenius',  'debit',  'TRUE'),
    ('fuel_payment',     'Bayar Bensin',        'jenius',  'debit',  'TRUE'),
    ('cash_in',          'Uang Cash Masuk',     'cash',    'credit', 'TRUE'),
    ('cash_out',         'Pengeluaran Cash',    'cash',    'debit',  'TRUE'),
    ('emoney_usage',     'Pemakaian eMoney',    'emoney',  'debit',  'TRUE'),
]
for row in seed_rows:
    add_table_row(seed_t, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 16 — SPESIFIKASI EMAIL
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '16. Spesifikasi Email', 1)
hr(doc)

add_heading(doc, '16.1 Email Magic Link Login', 2)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Atribut', 'Nilai'], bg='2D5FA8')
email_data = [
    ('Subject',    'Masuk ke Uang Kas — Link Login Anda'),
    ('From',       'noreply@uangkas.app (atau sesuai konfigurasi SMTP)'),
    ('To',         'email pengguna yang meminta login'),
    ('Berlaku',    '15 menit sejak dikirim'),
    ('Format',     'HTML + plaintext fallback'),
]
for row in email_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_para(doc, 'Isi Email (ringkasan konten):', bold=True)
add_bullet(doc, 'Salam: "Halo [nama pengguna],"')
add_bullet(doc, 'Instruksi: "Klik tombol di bawah untuk masuk ke aplikasi Uang Kas."')
add_bullet(doc, 'Tombol/Link: [Masuk ke Uang Kas] → URL: https://[domain]/auth/verify/[token]')
add_bullet(doc, 'Peringatan: "Link ini hanya berlaku selama 15 menit dan hanya dapat digunakan sekali."')
add_bullet(doc, 'Jika tidak merasa meminta: "Abaikan email ini. Akun Anda tetap aman."')
add_bullet(doc, 'Footer: nama aplikasi, tidak untuk dibalas (no-reply).')

doc.add_paragraph()
add_heading(doc, '16.2 Email Onboarding Driver', 2)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header(t, ['Atribut', 'Nilai'], bg='2D5FA8')
onboard_data = [
    ('Subject',    'Selamat Datang di Uang Kas — Akses Akun Anda'),
    ('From',       'noreply@uangkas.app'),
    ('To',         'email driver yang baru ditambahkan'),
    ('Berlaku',    '24 jam sejak dikirim'),
]
for row in onboard_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_para(doc, 'Isi Email (ringkasan konten):', bold=True)
add_bullet(doc, 'Salam: "Halo [nama driver],"')
add_bullet(doc, 'Penjelasan: "Anda telah didaftarkan sebagai driver di aplikasi Uang Kas oleh administrator."')
add_bullet(doc, 'Instruksi: "Klik tombol di bawah untuk mengaktifkan akun dan mulai menggunakan aplikasi."')
add_bullet(doc, 'Tombol/Link: [Aktifkan Akun Saya] → URL: https://[domain]/auth/verify/[token]')
add_bullet(doc, 'Peringatan: "Link ini hanya berlaku selama 24 jam."')
add_bullet(doc, 'Informasi: "Setelah login, Anda dapat mencatat transaksi dari smartphone Anda."')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 17 — PENANGANAN ERROR & VALIDASI
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '17. Penanganan Error & Validasi', 1)
hr(doc)

add_heading(doc, '17.1 HTTP Error Pages', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['HTTP Code', 'Kondisi', 'Tampilan'], bg='2D5FA8')
err_pages = [
    ('400 Bad Request',     'Input tidak valid, CSRF gagal',         'Halaman error dengan pesan "Permintaan tidak valid"'),
    ('403 Forbidden',       'Akses ke resource yang tidak diizinkan','Halaman error dengan pesan "Anda tidak memiliki akses"'),
    ('404 Not Found',       'URL tidak ditemukan',                   'Halaman error dengan link kembali ke dashboard'),
    ('429 Too Many Requests','Rate limit magic link terlampaui',     'Halaman error dengan info "Coba lagi dalam X menit"'),
    ('500 Internal Error',  'Unhandled exception di server',         'Halaman error generik, error detail di log server'),
]
for row in err_pages:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '17.2 Validasi Form — Aturan Umum', 2)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Tipe Field', 'Validasi Client-side', 'Validasi Server-side', 'Pesan Error'], bg='2D5FA8')
val_rules = [
    ('Email',     'HTML type=email, required',    'Regex RFC 5322, max 255 karakter, cek duplikat',   '"Format email tidak valid" / "Email sudah terdaftar"'),
    ('Nominal',   'HTML type=number, min=1000',   'Integer positif, tidak melebihi saldo jika debit', '"Nominal harus lebih dari Rp 1.000" / "Saldo tidak mencukupi (tersedia Rp X)"'),
    ('Tanggal',   'HTML type=date, max=today',    'Tidak di masa depan, format valid',                '"Tanggal tidak boleh di masa depan"'),
    ('Teks',      'maxlength di HTML',            'strip whitespace, max length sesuai kolom DB',     '"[Field] tidak boleh kosong" / "Melebihi [N] karakter"'),
    ('Select',    'required di HTML',             'Cek nilai ada di pilihan yang valid',              '"Pilihan tidak valid"'),
]
for row in val_rules:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '17.3 Flash Message — Kategori & Gaya', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Kategori', 'Warna', 'Contoh Pesan'], bg='2D5FA8')
flash_data = [
    ('success', 'Hijau (#28A745)',  '"Transaksi berhasil dicatat", "Driver berhasil ditambahkan"'),
    ('error',   'Merah (#DC3545)',  '"Saldo tidak mencukupi", "Email sudah terdaftar"'),
    ('warning', 'Kuning (#FFC107)', '"Session akan berakhir", "Link hanya berlaku 15 menit"'),
    ('info',    'Biru (#17A2B8)',   '"Email magic link telah dikirimkan ke [email]"'),
]
for row in flash_data:
    add_table_row(t, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 18 — SPESIFIKASI KEAMANAN TEKNIS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '18. Spesifikasi Keamanan Teknis', 1)
hr(doc)

add_heading(doc, '18.1 Implementasi CSRF Protection', 2)
add_bullet(doc, 'Menggunakan Flask-WTF CSRFProtect(app) yang diinisialisasi di extensions.py.')
add_bullet(doc, 'Setiap template HTML menyertakan {{ csrf_token() }} dalam hidden input di semua form POST.')
add_bullet(doc, 'Untuk request HTMX, tambahkan header: hx-headers=\'{"X-CSRFToken": "{{ csrf_token() }}"}\' di base template.')
add_bullet(doc, 'Jika CSRF token tidak valid: abort(400) dengan pesan "CSRF token tidak valid".')

doc.add_paragraph()
add_heading(doc, '18.2 Session Cookie Configuration', 2)
add_code(doc, "app.config['SESSION_COOKIE_HTTPONLY'] = True")
add_code(doc, "app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'")
add_code(doc, "app.config['SESSION_COOKIE_SECURE']   = True   # Hanya production (HTTPS)")
add_code(doc, "app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=365)  # 1 tahun")
add_code(doc, "app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY')  # WAJIB dari env var")

doc.add_paragraph()
add_heading(doc, '18.3 Hashing Magic Link Token', 2)
add_code(doc, "import secrets, hashlib")
add_code(doc, "")
add_code(doc, "# Generate token")
add_code(doc, "plain_token  = secrets.token_urlsafe(32)   # 32 bytes -> 43 char base64url")
add_code(doc, "hashed_token = hashlib.sha256(plain_token.encode()).hexdigest()")
add_code(doc, "")
add_code(doc, "# Simpan: hashed_token ke DB")
add_code(doc, "# Kirim: plain_token via email URL")
add_code(doc, "")
add_code(doc, "# Verifikasi saat klik link:")
add_code(doc, "input_hashed = hashlib.sha256(url_token.encode()).hexdigest()")
add_code(doc, "record = MagicToken.query.filter_by(token=input_hashed).first()")

doc.add_paragraph()
add_heading(doc, '18.4 Rate Limiting Magic Link', 2)
add_code(doc, "# Cek sebelum kirim magic link")
add_code(doc, "five_min_ago = datetime.utcnow() - timedelta(minutes=5)")
add_code(doc, "recent_count = MagicToken.query.filter(")
add_code(doc, "    MagicToken.user_id == user.id,")
add_code(doc, "    MagicToken.created_at >= five_min_ago")
add_code(doc, ").count()")
add_code(doc, "if recent_count >= 3:")
add_code(doc, "    abort(429)  # Too Many Requests")

doc.add_paragraph()
add_heading(doc, '18.5 Query Keamanan Data Driver', 2)
add_para(doc, 'Semua query yang mengambil data driver harus selalu disertai filter driver_id:')
add_code(doc, "# BENAR — selalu filter berdasarkan current_user")
add_code(doc, "wallets = Wallet.query.filter_by(driver_id=g.current_user.id).all()")
add_code(doc, "")
add_code(doc, "# SALAH — berpotensi akses data driver lain")
add_code(doc, "wallets = Wallet.query.filter_by(id=request.args.get('wallet_id')).first()")
add_code(doc, "")
add_code(doc, "# BENAR — validasi ownership sebelum akses by ID")
add_code(doc, "wallet = Wallet.query.filter_by(")
add_code(doc, "    id=wallet_id,")
add_code(doc, "    driver_id=g.current_user.id  # pastikan milik driver ini")
add_code(doc, ").first_or_404()")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 19 — SPESIFIKASI UI/UX
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '19. Spesifikasi UI/UX', 1)
hr(doc)

add_heading(doc, '19.1 Palet Warna', 2)
color_t = doc.add_table(rows=1, cols=4)
color_t.style = 'Table Grid'
add_table_header(color_t, ['Nama Warna', 'Hex', 'Digunakan Untuk', 'Catatan'], bg='2D5FA8')
colors = [
    ('Primary',        '#1E3A5F', 'Heading, navbar background, teks utama brand',  'Navy gelap'),
    ('Primary Light',  '#2D5FA8', 'Tombol primer, link aktif, header tabel',        'Biru sedang'),
    ('Accent',         '#3B82F6', 'Hover state, focus ring, highlight',             'Biru cerah'),
    ('Success',        '#28A745', 'Flash success, badge aktif, kredit (+)',          'Hijau'),
    ('Danger',         '#DC3545', 'Flash error, badge nonaktif, debit (-)',          'Merah'),
    ('Warning',        '#FFC107', 'Flash warning, saldo rendah',                    'Kuning'),
    ('Info',           '#17A2B8', 'Flash info, tooltip',                            'Cyan'),
    ('Background',     '#F8FAFF', 'Background halaman utama',                       'Biru sangat muda'),
    ('Card',           '#FFFFFF', 'Background kartu/tabel',                         'Putih'),
    ('Jenius Card',    '#1E3A5F', 'Kartu wallet Jenius di dashboard driver',        'Navy'),
    ('eMoney Card',    '#1A7A5A', 'Kartu wallet eMoney di dashboard driver',        'Hijau tua'),
    ('Cash Card',      '#C45C00', 'Kartu wallet Cash di dashboard driver',          'Oranye tua'),
]
for row in colors:
    add_table_row(color_t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '19.2 Tipografi', 2)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Elemen', 'Font', 'Ukuran', 'Weight'], bg='2D5FA8')
typo_data = [
    ('Judul H1',            'Plus Jakarta Sans', '28px', 'Bold (700)'),
    ('Judul H2',            'Plus Jakarta Sans', '22px', 'SemiBold (600)'),
    ('Judul H3',            'Plus Jakarta Sans', '18px', 'SemiBold (600)'),
    ('Body teks',           'Plus Jakarta Sans', '14px', 'Regular (400)'),
    ('Label form',          'Plus Jakarta Sans', '13px', 'Medium (500)'),
    ('Saldo (angka besar)', 'JetBrains Mono',    '24px', 'Bold (700)'),
    ('Nominal transaksi',   'JetBrains Mono',    '14px', 'Regular (400)'),
    ('Kode/snippet',        'JetBrains Mono',    '12px', 'Regular (400)'),
]
for row in typo_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '19.3 Breakpoint Responsif', 2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Breakpoint', 'Lebar Layar', 'Penyesuaian Layout'], bg='2D5FA8')
bp_data = [
    ('Mobile',  '< 640px',   'Satu kolom, navigasi hamburger, tombol full-width, tabel scroll horizontal'),
    ('Tablet',  '640–1024px','Dua kolom untuk kartu wallet, navigasi sidebar collapsed'),
    ('Desktop', '> 1024px',  'Sidebar navigasi tampil penuh, tabel dengan lebih banyak kolom'),
]
for row in bp_data:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_heading(doc, '19.4 Format Tampilan Rupiah', 2)
add_para(doc, 'Semua nominal rupiah ditampilkan dengan format: Rp [angka dengan pemisah ribuan]')
add_code(doc, "# Python Jinja2 filter")
add_code(doc, "def format_rupiah(value):")
add_code(doc, '    return f"Rp {value:,.0f}".replace(",", ".")')
add_code(doc, "")
add_code(doc, "# Contoh output:")
add_code(doc, "# 1500000  ->  Rp 1.500.000")
add_code(doc, "# 250000   ->  Rp 250.000")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 20 — SKENARIO PENGUJIAN FUNGSIONAL
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '20. Skenario Pengujian Fungsional', 1)
hr(doc)

add_para(doc, 'Skenario pengujian berikut mencakup happy path dan edge case untuk modul utama. '
         'Setiap skenario diuji manual oleh QA sebelum go-live.', italic=True)

doc.add_paragraph()

test_groups = [
    ('Autentikasi', [
        ('TC-AUTH-01', 'Login owner dengan email valid',
         'Email setyanaputra@yahoo.com di form login',
         'Email magic link terkirim. Halaman konfirmasi tampil.'),
        ('TC-AUTH-02', 'Verifikasi magic link valid',
         'Klik magic link yang baru diterima di email (dalam 15 menit)',
         'Redirect ke /owner/. Session aktif. Cookie tersimpan.'),
        ('TC-AUTH-03', 'Magic link kedaluwarsa (>15 menit)',
         'Klik magic link setelah 16 menit',
         'Redirect ke /auth/login. Flash "Link tidak valid atau sudah kedaluwarsa".'),
        ('TC-AUTH-04', 'Magic link digunakan dua kali',
         'Klik magic link yang sama dua kali (klik kedua setelah klik pertama berhasil)',
         'Klik pertama: login berhasil. Klik kedua: redirect ke login dengan pesan error.'),
        ('TC-AUTH-05', 'Rate limit: 3 request dalam 5 menit',
         'Request magic link 4 kali dalam 5 menit',
         'Tiga pertama: email terkirim. Keempat: halaman 429 Too Many Requests.'),
        ('TC-AUTH-06', 'Logout',
         'Klik tombol logout saat sudah login',
         'Cookie dihapus. Redirect ke /auth/login. Session di DB is_active=FALSE.'),
        ('TC-AUTH-07', 'Akses driver ke halaman owner',
         'Driver yang login mencoba akses URL /owner/',
         'Redirect ke /driver/ atau tampilan 403 Forbidden.'),
    ]),
    ('Manajemen Driver (Owner)', [
        ('TC-DRV-01', 'Tambah driver baru',
         'Form tambah driver: nama "Budi Santoso", email "budi@email.com"',
         'Driver tersimpan. 3 wallet dibuat (saldo 0). Email onboarding terkirim.'),
        ('TC-DRV-02', 'Tambah driver dengan email duplikat',
         'Email yang sudah terdaftar di form tambah driver',
         'Form ditolak. Flash error "Email sudah terdaftar".'),
        ('TC-DRV-03', 'Nonaktifkan driver',
         'Toggle nonaktif pada driver aktif',
         'is_active=FALSE. Session driver diterminasi. Driver tidak bisa login.'),
        ('TC-DRV-04', 'Login driver yang dinonaktifkan',
         'Driver nonaktif coba login dengan email mereka',
         'Magic link terkirim (tidak membocorkan status), tapi saat klik: "Akun tidak aktif".'),
    ]),
    ('Transaksi Jenius — Tarik ATM', [
        ('TC-JNS-01', 'Tarik ATM normal',
         'Driver: Tarik ATM Rp 200.000, saldo Jenius Rp 500.000',
         'Jenius berkurang 200.000. Cash bertambah 200.000. Dua transaksi terhubung (linked_txn_id).'),
        ('TC-JNS-02', 'Tarik ATM melebihi saldo',
         'Driver: Tarik ATM Rp 600.000, saldo Jenius Rp 500.000',
         'Transaksi ditolak. Flash error "Saldo Jenius tidak mencukupi (tersedia Rp 500.000)".'),
        ('TC-JNS-03', 'Bayar bensin',
         'Driver: Bayar Bensin Rp 150.000, saldo Jenius Rp 500.000',
         'Jenius berkurang 150.000. Tidak ada perubahan pada Cash. Satu transaksi.'),
        ('TC-JNS-04', 'Atomicity linked transaction',
         'Simulasi error saat INSERT Cash (force error di kode)',
         'Kedua transaksi di-rollback. Saldo Jenius dan Cash tidak berubah.'),
    ]),
    ('Transaksi Cash', [
        ('TC-CSH-01', 'Catat pengeluaran cash normal',
         'Driver: Rp 15.000, catatan "Parkir Sudirman", saldo Cash Rp 200.000',
         'Cash berkurang 15.000. Transaksi tersimpan dengan catatan.'),
        ('TC-CSH-02', 'Pengeluaran cash melebihi saldo',
         'Driver: Rp 300.000, saldo Cash Rp 200.000',
         'Transaksi ditolak. Flash error "Saldo cash tidak mencukupi".'),
        ('TC-CSH-03', 'Catatan wajib diisi',
         'Submit form pengeluaran cash tanpa catatan',
         'Form ditolak client-side (required) dan server-side. Flash error.'),
    ]),
    ('Manajemen Jenis Transaksi', [
        ('TC-TXN-01', 'Tambah jenis transaksi baru',
         'Owner: label "Bayar Cuci Mobil", wallet Jenius, arah debit',
         'Jenis transaksi tersimpan. Muncul di dropdown driver saat catat transaksi Jenius.'),
        ('TC-TXN-02', 'Hapus jenis transaksi is_system=TRUE',
         'Owner coba akses delete URL untuk jenis transaksi system',
         'abort(403). Tombol hapus tidak tampil di UI untuk is_system=TRUE.'),
        ('TC-TXN-03', 'Hapus jenis transaksi yang sudah dipakai',
         'Owner coba hapus jenis transaksi non-system yang sudah dipakai 1 transaksi',
         'Error "Tidak dapat menghapus, sudah digunakan di 1 transaksi".'),
    ]),
    ('Session Management', [
        ('TC-SES-01', 'Owner terminate session driver',
         'Owner klik terminate pada session driver aktif',
         'Session di DB is_active=FALSE. Driver redirect ke login pada request berikutnya.'),
        ('TC-SES-02', 'Driver akses setelah session diterminasi',
         'Driver yang session-nya diterminasi membuat request apapun',
         'Cookie dihapus. Redirect ke /auth/login. Flash "Session tidak valid".'),
    ]),
]

for group_name, tests in test_groups:
    add_heading(doc, f'20.x {group_name}', 2)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    add_table_header(t, ['ID Test', 'Skenario', 'Input / Kondisi', 'Hasil yang Diharapkan'])
    for tc in tests:
        add_table_row(t, list(tc), font_size=8.5)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PENUTUP
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Persetujuan Dokumen', 1)
hr(doc)
add_para(doc, 'Dokumen ini telah ditinjau dan disetujui oleh pihak-pihak berikut:')
doc.add_paragraph()

sign_t = doc.add_table(rows=4, cols=3)
sign_t.style = 'Table Grid'
add_table_header(sign_t, ['Peran', 'Nama', 'Tanda Tangan & Tanggal'])
sign_data = [
    ('Owner / Product Owner',   'Setyana Putra',  ' '),
    ('Lead Developer',          ' ',              ' '),
    ('Quality Assurance',       ' ',              ' '),
]
for row_data in sign_data:
    row = sign_t.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        if cell.text.strip():
            cell.paragraphs[0].runs[0].font.size = Pt(10)

# ─── Save ────────────────────────────────────────────────────────────────────

output_path = r'c:\Python312\flask_uangkas\FSD_UangKas_v1.0.docx'
doc.save(output_path)
print(f"FSD berhasil dibuat: {output_path}")
