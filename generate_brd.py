"""
Generate BRD (Business Requirement Document) for Uang Kas v1.0
Run: python generate_brd.py
Output: BRD_UangKas_v1.0.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', '2D5FA8'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_table_header(table, headers, bg='1E3A5F', fg='FFFFFF', font_size=9):
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(fg)

def add_table_row(table, values, bg=None, bold=False, align_center_cols=None, font_size=9):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        if align_center_cols and i in align_center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(font_size)
    return row

def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return para

def add_para(doc, text, bold=False, italic=False, size=10, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2D5FA8')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Document Setup ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('BUSINESS REQUIREMENT DOCUMENT')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()

cover_app = doc.add_paragraph()
cover_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_app.add_run('UANG KAS')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x2D, 0x5F, 0xA8)

cover_ver = doc.add_paragraph()
cover_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_ver.add_run('Versi 1.0')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

# Info box cover
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'
info_data = [
    ('Nama Aplikasi',   'Uang Kas'),
    ('Versi Dokumen',   '1.0'),
    ('Tanggal',         '16 April 2026'),
    ('Dibuat Oleh',     'Tim Pengembang'),
    ('Status',          'Draft'),
]
for i, (k, v) in enumerate(info_data):
    row = info_table.rows[i]
    set_cell_bg(row.cells[0], 'DCE6F1')
    row.cells[0].text = k
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = v
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

for row in info_table.rows:
    for cell in row.cells:
        cell.width = Cm(6)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DOCUMENT INFORMATION
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Informasi Dokumen', 1)
hr(doc)

add_heading(doc, '1.1 Riwayat Revisi', 2)
rev_table = doc.add_table(rows=1, cols=5)
rev_table.style = 'Table Grid'
add_table_header(rev_table, ['Versi', 'Tanggal', 'Deskripsi', 'Dibuat Oleh', 'Disetujui Oleh'])
add_table_row(rev_table, ['1.0', '16 Apr 2026', 'Initial Draft — BRD Uang Kas', 'Tim Dev', 'Owner'],
              bg='EBF3FF', align_center_cols=[0,1])

doc.add_paragraph()

add_heading(doc, '1.2 Daftar Istilah', 2)
term_table = doc.add_table(rows=1, cols=2)
term_table.style = 'Table Grid'
add_table_header(term_table, ['Istilah / Singkatan', 'Definisi'])
terms = [
    ('Owner',           'Pemilik bisnis, memiliki akses penuh sebagai administrator aplikasi'),
    ('Driver',          'Supir/pengemudi yang menggunakan wallet untuk operasional harian'),
    ('Jenius',          'Kartu ATM/debit bank Jenius yang dipegang driver'),
    ('eMoney',          'Kartu uang elektronik untuk pembayaran tol dan parkir'),
    ('Uang Cash',       'Uang tunai hasil penarikan ATM yang dikelola driver'),
    ('Magic Link',      'Tautan autentikasi satu kali yang dikirim ke email pengguna'),
    ('Wallet',          'Dompet digital dalam sistem yang mewakili saldo kartu/uang'),
    ('Top Up',          'Penambahan saldo oleh owner ke wallet driver'),
    ('Linked Txn',      'Pasangan transaksi yang saling terhubung, contoh: tarik ATM → cash masuk'),
    ('BRD',             'Business Requirement Document — dokumen spesifikasi kebutuhan bisnis'),
    ('MPA',             'Multi-Page Application — arsitektur web dengan halaman terpisah per fungsi'),
    ('HTMX',            'Library JavaScript ringan untuk partial page update tanpa full reload'),
    ('ORM',             'Object Relational Mapper — abstraksi database via kode Python (SQLAlchemy)'),
    ('IDR',             'Indonesian Rupiah — satuan mata uang, disimpan sebagai integer'),
    ('ATM',             'Anjungan Tunai Mandiri — mesin untuk penarikan uang tunai'),
    ('SPBU',            'Stasiun Pengisian Bahan Bakar Umum / Pom Bensin'),
]
for term, defn in terms:
    add_table_row(term_table, [term, defn], font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Ringkasan Eksekutif', 1)
hr(doc)

add_para(doc, 'Uang Kas adalah aplikasi web monitoring keuangan operasional yang dirancang untuk membantu owner '
         'dalam memantau penggunaan dana yang dipegang oleh dua orang driver secara real-time. Saat ini pencatatan '
         'dilakukan secara manual sehingga rentan terhadap kesalahan, keterlambatan pelaporan, dan kurangnya transparansi.')

add_para(doc, 'Setiap driver memegang tiga instrumen keuangan: (1) Kartu ATM Jenius untuk penarikan tunai dan '
         'pembayaran bahan bakar, (2) Kartu eMoney untuk pembayaran tol dan parkir, serta (3) Uang Cash hasil '
         'penarikan ATM yang digunakan untuk keperluan harian seperti parkir dan tip.')

add_para(doc, 'Aplikasi ini menyediakan dashboard monitoring terpusat bagi owner serta antarmuka pencatatan '
         'transaksi yang sederhana dan mobile-friendly bagi driver. Autentikasi menggunakan magic link tanpa '
         'password untuk kemudahan penggunaan di perangkat mobile.')

doc.add_paragraph()
add_heading(doc, '2.1 Tujuan Bisnis', 2)
goals = [
    'Memberikan visibilitas real-time kepada owner atas seluruh saldo dan transaksi yang dilakukan driver.',
    'Mengurangi risiko penyalahgunaan dana melalui pencatatan terstruktur dan audit trail yang lengkap.',
    'Menyederhanakan proses pelaporan keuangan operasional harian.',
    'Menghilangkan ketergantungan pada pencatatan manual (kertas / WhatsApp).',
    'Memudahkan driver dalam mencatat pengeluaran melalui antarmuka mobile yang intuitif.',
]
for g in goals:
    add_bullet(doc, g)

doc.add_paragraph()
add_heading(doc, '2.2 Ruang Lingkup', 2)

add_para(doc, 'DALAM RUANG LINGKUP (In Scope):', bold=True)
in_scope = [
    'Manajemen pengguna: owner dan driver (maksimal konfigurasi awal 2 driver)',
    'Monitoring tiga jenis wallet per driver: Jenius, eMoney, dan Uang Cash',
    'Pencatatan transaksi oleh driver (tarik ATM, bayar bensin, pengeluaran cash)',
    'Pengelolaan saldo oleh owner (top up, penarikan Jenius; top up eMoney)',
    'Laporan dan riwayat transaksi per driver per wallet',
    'Autentikasi magic link berbasis email',
    'Manajemen session driver oleh owner (terminate session)',
    'Konfigurasi jenis transaksi oleh owner (tambah/edit transaction types)',
]
for item in in_scope:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, 'DI LUAR RUANG LINGKUP (Out of Scope):', bold=True)
out_scope = [
    'Integrasi langsung dengan API bank Jenius atau sistem eMoney',
    'Fitur notifikasi push (push notification) ke perangkat mobile',
    'Laporan pajak atau akuntansi formal',
    'Multi-owner atau struktur organisasi hierarkis',
    'Aplikasi native mobile (Android/iOS)',
    'Fitur penggajian atau slip gaji driver',
]
for item in out_scope:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — STAKEHOLDERS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Pemangku Kepentingan', 1)
hr(doc)

sh_table = doc.add_table(rows=1, cols=4)
sh_table.style = 'Table Grid'
add_table_header(sh_table, ['Peran', 'Nama/Email', 'Tanggung Jawab', 'Tingkat Keterlibatan'])
sh_data = [
    ('Owner / Admin',  'setyanaputra@yahoo.com',
     'Pemilik bisnis; menyetujui requirement; menggunakan fitur admin dan monitoring',
     'Tinggi — Primary User'),
    ('Driver 1',       'Terdaftar oleh Owner',
     'Mencatat transaksi harian; menggunakan antarmuka mobile driver',
     'Tinggi — Primary User'),
    ('Driver 2',       'Terdaftar oleh Owner',
     'Mencatat transaksi harian; menggunakan antarmuka mobile driver',
     'Tinggi — Primary User'),
    ('Tim Pengembang', '-',
     'Merancang, membangun, dan memelihara aplikasi',
     'Tinggi — Developer'),
]
for row_data in sh_data:
    add_table_row(sh_table, list(row_data), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — FUNCTIONAL REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Kebutuhan Fungsional', 1)
hr(doc)

# ── 4.1 Autentikasi ──────────────────────────────────────────────────────────
add_heading(doc, '4.1 Autentikasi & Session', 2)

add_para(doc, 'FR-AUTH-01: Magic Link Login', bold=True)
add_bullet(doc, 'Pengguna (owner atau driver) memasukkan alamat email di halaman login.')
add_bullet(doc, 'Sistem memvalidasi apakah email terdaftar dan akun aktif.')
add_bullet(doc, 'Jika valid, sistem mengirimkan magic link ke email pengguna.')
add_bullet(doc, 'Magic link berisi token unik yang di-hash sebelum disimpan di database.')
add_bullet(doc, 'Token berlaku selama 15 menit dan hanya dapat digunakan satu kali.')
add_bullet(doc, 'Setelah verifikasi berhasil, sistem membuat session aktif dan mengarahkan pengguna ke dashboard sesuai peran.')

doc.add_paragraph()
add_para(doc, 'FR-AUTH-02: Session Management', bold=True)
add_bullet(doc, 'Session owner: tidak ada waktu kedaluwarsa otomatis (persistent).')
add_bullet(doc, 'Session driver: tidak ada waktu kedaluwarsa otomatis, namun dapat diterminasi oleh owner.')
add_bullet(doc, 'Owner dapat melihat daftar session driver yang aktif beserta informasi perangkat dan waktu login.')
add_bullet(doc, 'Owner dapat menterminasi session driver tertentu. Driver yang session-nya diterminasi akan otomatis logout pada request berikutnya.')
add_bullet(doc, 'Pengguna dapat melakukan logout mandiri (self-logout).')

doc.add_paragraph()
add_para(doc, 'FR-AUTH-03: Onboarding Driver (Magic Link Registrasi)', bold=True)
add_bullet(doc, 'Ketika owner menambahkan driver baru, sistem mengirimkan magic link onboarding ke email driver.')
add_bullet(doc, 'Magic link onboarding memiliki masa berlaku 24 jam.')
add_bullet(doc, 'Driver yang mengklik magic link onboarding langsung masuk ke aplikasi dengan status akun aktif.')

doc.add_paragraph()

# ── 4.2 Manajemen Pengguna ───────────────────────────────────────────────────
add_heading(doc, '4.2 Manajemen Pengguna (Owner)', 2)

add_para(doc, 'FR-USER-01: Tambah Driver', bold=True)
add_bullet(doc, 'Owner dapat menambahkan driver baru dengan mengisi: nama lengkap dan alamat email.')
add_bullet(doc, 'Sistem otomatis membuat tiga wallet untuk driver baru: Jenius, eMoney, dan Uang Cash dengan saldo awal 0.')
add_bullet(doc, 'Sistem mengirimkan magic link onboarding ke email driver.')

doc.add_paragraph()
add_para(doc, 'FR-USER-02: Edit & Nonaktifkan Driver', bold=True)
add_bullet(doc, 'Owner dapat mengubah nama driver.')
add_bullet(doc, 'Owner dapat menonaktifkan akses driver (is_active = FALSE). Driver yang dinonaktifkan tidak dapat login.')
add_bullet(doc, 'Owner dapat mengaktifkan kembali akun driver yang dinonaktifkan.')

doc.add_paragraph()
add_para(doc, 'FR-USER-03: Kirim Ulang Magic Link', bold=True)
add_bullet(doc, 'Owner dapat meminta pengiriman ulang magic link onboarding untuk driver yang belum login.')

doc.add_paragraph()

# ── 4.3 Manajemen Wallet ─────────────────────────────────────────────────────
add_heading(doc, '4.3 Manajemen Wallet (Owner)', 2)

add_para(doc, 'FR-WALLET-01: Lihat Semua Wallet', bold=True)
add_bullet(doc, 'Owner dapat melihat daftar semua wallet seluruh driver beserta saldo terkini.')
add_bullet(doc, 'Tampilan dikelompokkan per driver.')

doc.add_paragraph()
add_para(doc, 'FR-WALLET-02: Top Up Jenius', bold=True)
add_bullet(doc, 'Owner dapat menambahkan saldo ke wallet Jenius milik driver tertentu.')
add_bullet(doc, 'Input yang diperlukan: nominal (Rupiah), tanggal transaksi, dan catatan opsional.')
add_bullet(doc, 'Sistem mencatat transaksi dengan tipe owner_topup dan memperbarui saldo wallet.')

doc.add_paragraph()
add_para(doc, 'FR-WALLET-03: Penarikan Jenius (Owner)', bold=True)
add_bullet(doc, 'Owner dapat mengurangi saldo wallet Jenius (fitur jarang digunakan, namun tersedia).')
add_bullet(doc, 'Input: nominal, tanggal, catatan.')
add_bullet(doc, 'Sistem mencatat transaksi dengan tipe owner_withdrawal.')

doc.add_paragraph()
add_para(doc, 'FR-WALLET-04: Top Up eMoney', bold=True)
add_bullet(doc, 'Owner dapat menambahkan saldo ke wallet eMoney milik driver tertentu.')
add_bullet(doc, 'Input: nominal, tanggal, catatan opsional.')

doc.add_paragraph()
add_para(doc, 'FR-WALLET-05: Tambah Wallet Baru', bold=True)
add_bullet(doc, 'Owner dapat menambahkan wallet baru untuk driver jika di masa depan ada instrumen keuangan tambahan.')
add_bullet(doc, 'Input: nama wallet, jenis (label bebas), dan driver yang bersangkutan.')

doc.add_paragraph()

# ── 4.4 Manajemen Jenis Transaksi ───────────────────────────────────────────
add_heading(doc, '4.4 Manajemen Jenis Transaksi (Owner)', 2)

add_para(doc, 'FR-TXNTYPE-01: Lihat Jenis Transaksi', bold=True)
add_bullet(doc, 'Owner dapat melihat seluruh daftar jenis transaksi yang terdaftar di sistem.')
add_bullet(doc, 'Kolom yang ditampilkan: kode, label, berlaku untuk wallet, arah (kredit/debit), status aktif, is_system.')

doc.add_paragraph()
add_para(doc, 'FR-TXNTYPE-02: Tambah Jenis Transaksi', bold=True)
add_bullet(doc, 'Owner dapat menambahkan jenis transaksi baru.')
add_bullet(doc, 'Input: label tampilan, wallet yang berlaku, arah default (kredit/debit), urutan tampil.')
add_bullet(doc, 'Sistem otomatis membuat kode unik dari label.')
add_bullet(doc, 'Jenis transaksi baru memiliki is_system = FALSE.')

doc.add_paragraph()
add_para(doc, 'FR-TXNTYPE-03: Edit Jenis Transaksi', bold=True)
add_bullet(doc, 'Owner dapat mengedit label dan urutan tampil untuk semua jenis transaksi.')
add_bullet(doc, 'Jenis transaksi dengan is_system = TRUE: tidak dapat dihapus, tidak dapat diubah kode dan arahnya.')
add_bullet(doc, 'Jenis transaksi dengan is_system = FALSE: dapat dinonaktifkan atau dihapus selama belum digunakan di transaksi manapun.')

doc.add_paragraph()

# ── 4.5 Transaksi Driver ────────────────────────────────────────────────────
add_heading(doc, '4.5 Pencatatan Transaksi (Driver)', 2)

add_para(doc, 'FR-TXN-01: Catat Transaksi Jenius', bold=True)
add_bullet(doc, 'Driver dapat membuka halaman "Catat Transaksi Jenius" dari menu.')
add_bullet(doc, 'Driver memilih jenis transaksi dari pilihan yang tersedia:')
add_bullet(doc, '    • Tarik ATM — sistem secara otomatis membuat linked transaction: debit Jenius + kredit Uang Cash dengan nominal yang sama.')
add_bullet(doc, '    • Bayar Bensin — debit Jenius saja.')
add_bullet(doc, '    • Transaksi Lainnya — driver memilih dari dropdown jenis transaksi aktif yang berlaku untuk wallet Jenius.')
add_bullet(doc, 'Input wajib: jenis transaksi, nominal, tanggal & jam transaksi.')
add_bullet(doc, 'Input opsional: catatan/keterangan.')
add_bullet(doc, 'Sistem memvalidasi saldo Jenius mencukupi sebelum menyimpan transaksi debit.')

doc.add_paragraph()
add_para(doc, 'FR-TXN-02: Catat Pengeluaran Uang Cash', bold=True)
add_bullet(doc, 'Driver dapat membuka halaman "Catat Pengeluaran Cash" dari menu.')
add_bullet(doc, 'Input wajib: nominal, tanggal & jam, catatan (untuk apa uang digunakan).')
add_bullet(doc, 'Sistem mencatat transaksi dengan tipe cash_out dan mengurangi saldo wallet Cash.')
add_bullet(doc, 'Sistem memvalidasi saldo Cash mencukupi sebelum menyimpan.')

doc.add_paragraph()

# ── 4.6 Monitoring ──────────────────────────────────────────────────────────
add_heading(doc, '4.6 Monitoring & Laporan', 2)

add_para(doc, 'FR-MON-01: Dashboard Owner', bold=True)
add_bullet(doc, 'Menampilkan ringkasan saldo semua wallet semua driver dalam satu halaman.')
add_bullet(doc, 'Menampilkan total saldo keseluruhan (Jenius + eMoney + Cash) per driver.')
add_bullet(doc, 'Menampilkan 5 transaksi terakhir dari semua driver.')

doc.add_paragraph()
add_para(doc, 'FR-MON-02: Dashboard Driver', bold=True)
add_bullet(doc, 'Menampilkan tiga kartu saldo milik driver yang login: Jenius, eMoney, dan Uang Cash.')
add_bullet(doc, 'Menampilkan 5 transaksi terakhir milik driver tersebut.')

doc.add_paragraph()
add_para(doc, 'FR-MON-03: Laporan Transaksi Jenius (Owner)', bold=True)
add_bullet(doc, 'Owner dapat memilih driver dan rentang tanggal untuk melihat riwayat transaksi Jenius.')
add_bullet(doc, 'Kolom: tanggal/jam, jenis transaksi, nominal, saldo sebelum, saldo sesudah, dicatat oleh, catatan.')
add_bullet(doc, 'Tersedia filter: per driver, per jenis transaksi, rentang tanggal.')

doc.add_paragraph()
add_para(doc, 'FR-MON-04: Laporan Uang Cash (Owner)', bold=True)
add_bullet(doc, 'Owner dapat melihat laporan cash per driver dengan informasi: saldo awal periode, daftar transaksi masuk (dari ATM), daftar transaksi keluar (pengeluaran), dan saldo akhir periode.')
add_bullet(doc, 'Tersedia filter rentang tanggal.')

doc.add_paragraph()
add_para(doc, 'FR-MON-05: Riwayat Transaksi Driver', bold=True)
add_bullet(doc, 'Driver dapat melihat riwayat transaksi Jenius dan Cash miliknya sendiri.')
add_bullet(doc, 'Tampilan berurutan dari terbaru ke terlama.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — NON-FUNCTIONAL REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Kebutuhan Non-Fungsional', 1)
hr(doc)

nfr_table = doc.add_table(rows=1, cols=4)
nfr_table.style = 'Table Grid'
add_table_header(nfr_table, ['ID', 'Kategori', 'Kebutuhan', 'Kriteria Penerimaan'])

nfr_data = [
    ('NFR-01', 'Performa',
     'Halaman dashboard harus termuat dalam waktu wajar',
     'Response time < 2 detik pada koneksi 4G mobile'),
    ('NFR-02', 'Ketersediaan',
     'Aplikasi harus tersedia untuk penggunaan harian',
     'Uptime ≥ 99% pada jam operasional (06.00–22.00 WIB)'),
    ('NFR-03', 'Keamanan',
     'Session cookie harus aman dari intercept',
     'Cookie menggunakan flag HTTPONLY dan SAMESITE=Lax'),
    ('NFR-04', 'Keamanan',
     'Token magic link harus disimpan dalam bentuk hash',
     'Token di database di-hash menggunakan SHA-256; token plaintext hanya dikirim via email'),
    ('NFR-05', 'Keamanan',
     'Proteksi CSRF pada semua form POST',
     'Setiap form POST menyertakan CSRF token yang divalidasi server'),
    ('NFR-06', 'Keamanan',
     'Akses berbasis peran (Role-Based Access Control)',
     'Endpoint owner tidak dapat diakses oleh driver dan sebaliknya; redirect ke 403 jika melanggar'),
    ('NFR-07', 'Usability',
     'Antarmuka driver harus ramah mobile',
     'Tampilan responsif, tombol aksi berukuran minimal 44px, mudah digunakan satu tangan'),
    ('NFR-08', 'Maintainability',
     'Kode terstruktur dan mudah dipelihara',
     'Menggunakan Blueprint Flask; model terpisah per entitas; tidak ada logika bisnis di template'),
    ('NFR-09', 'Portability',
     'Database kompatibel untuk development dan production',
     'SQLite untuk development; PostgreSQL untuk production; tidak menggunakan fitur DB-spesifik'),
    ('NFR-10', 'Akurasi Data',
     'Saldo tidak boleh bernilai negatif',
     'Sistem menolak transaksi debit jika saldo tidak mencukupi dengan pesan error yang jelas'),
    ('NFR-11', 'Audit Trail',
     'Semua transaksi harus tercatat lengkap',
     'Setiap transaksi menyimpan: balance_before, balance_after, created_by, created_at'),
    ('NFR-12', 'Integritas Data',
     'Transaksi linked (ATM + Cash) harus atomic',
     'Jika salah satu INSERT gagal, keduanya di-rollback; menggunakan database transaction'),
]
for row_data in nfr_data:
    add_table_row(nfr_table, list(row_data), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DATA REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Kebutuhan Data — Skema Database', 1)
hr(doc)

add_para(doc, 'Semua tabel menggunakan UUID sebagai primary key. Saldo dan nominal transaksi disimpan '
         'sebagai BIGINT (integer Rupiah) untuk menghindari floating-point error.', italic=True)

doc.add_paragraph()

# ── Tabel users ──────────────────────────────────────────────────────────────
add_heading(doc, '6.1 Tabel: users', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
users_cols = [
    ('id',         'UUID',         'PK',              'gen_random_uuid()', 'Primary key'),
    ('email',      'VARCHAR(255)', 'UNIQUE, NOT NULL', '-',                'Alamat email unik'),
    ('name',       'VARCHAR(100)', 'NOT NULL',         '-',                'Nama lengkap'),
    ('role',       'VARCHAR(10)',  'NOT NULL',         '-',                "Nilai: 'owner' atau 'driver'"),
    ('is_active',  'BOOLEAN',      'NOT NULL',         'TRUE',             'FALSE = akun dinonaktifkan'),
    ('created_at', 'TIMESTAMP',    'NOT NULL',         'NOW()',            'Waktu pembuatan akun'),
    ('updated_at', 'TIMESTAMP',    'NOT NULL',         'NOW()',            'Waktu update terakhir'),
]
for row in users_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()

# ── Tabel magic_tokens ───────────────────────────────────────────────────────
add_heading(doc, '6.2 Tabel: magic_tokens', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
mt_cols = [
    ('id',         'UUID',         'PK',              'gen_random_uuid()', 'Primary key'),
    ('user_id',    'UUID',         'FK → users, NOT NULL', '-',            'Pemilik token'),
    ('token',      'VARCHAR(128)', 'UNIQUE, NOT NULL', '-',                'Token di-hash SHA-256'),
    ('purpose',    'VARCHAR(20)',  'NOT NULL',         '-',                "Nilai: 'login' atau 'onboarding'"),
    ('expires_at', 'TIMESTAMP',    'NOT NULL',         '-',                'Login: +15 menit; Onboarding: +24 jam'),
    ('used_at',    'TIMESTAMP',    'NULL',             'NULL',             'NULL = belum digunakan'),
    ('created_at', 'TIMESTAMP',    'NOT NULL',         'NOW()',            'Waktu pembuatan token'),
]
for row in mt_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()

# ── Tabel user_sessions ──────────────────────────────────────────────────────
add_heading(doc, '6.3 Tabel: user_sessions', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
us_cols = [
    ('id',              'UUID',         'PK',                   'gen_random_uuid()', 'Primary key'),
    ('user_id',         'UUID',         'FK → users, NOT NULL', '-',                 'Pemilik session'),
    ('session_token',   'VARCHAR(255)', 'UNIQUE, NOT NULL',     '-',                 'Token disimpan di Flask cookie'),
    ('device_info',     'VARCHAR(255)', 'NULL',                 'NULL',              'User-Agent browser/app'),
    ('ip_address',      'VARCHAR(45)',  'NULL',                 'NULL',              'IP address saat login (IPv4/IPv6)'),
    ('is_active',       'BOOLEAN',      'NOT NULL',             'TRUE',              'FALSE = session sudah diterminasi'),
    ('created_at',      'TIMESTAMP',    'NOT NULL',             'NOW()',             'Waktu login'),
    ('terminated_at',   'TIMESTAMP',    'NULL',                 'NULL',              'Waktu terminasi (jika diterminasi)'),
    ('terminated_by',   'UUID',         'FK → users, NULL',     'NULL',              'Owner yang menterminasi; NULL = self-logout'),
]
for row in us_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()

# ── Tabel transaction_types ──────────────────────────────────────────────────
add_heading(doc, '6.4 Tabel: transaction_types', 2)
add_para(doc, 'Menggantikan ENUM txn_type agar owner dapat menambah/mengedit jenis transaksi secara dinamis.', italic=True)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
tt_cols = [
    ('id',           'UUID',         'PK',                   'gen_random_uuid()', 'Primary key'),
    ('code',         'VARCHAR(50)',  'UNIQUE, NOT NULL',     '-',                 'Kode unik, e.g. atm_withdrawal'),
    ('label',        'VARCHAR(100)', 'NOT NULL',             '-',                 'Label tampilan, e.g. "Tarik ATM"'),
    ('wallet_type',  'VARCHAR(20)',  'NOT NULL',             '-',                 "Nilai: 'jenius', 'emoney', 'cash', 'all'"),
    ('direction',    'VARCHAR(10)',  'NOT NULL',             '-',                 "Nilai: 'credit', 'debit', 'both'"),
    ('is_system',    'BOOLEAN',      'NOT NULL',             'FALSE',             'TRUE = data seed, tidak bisa dihapus'),
    ('is_active',    'BOOLEAN',      'NOT NULL',             'TRUE',              'FALSE = tidak muncul di pilihan driver'),
    ('sort_order',   'SMALLINT',     'NOT NULL',             '99',                'Urutan tampil di dropdown'),
    ('created_by',   'UUID',         'FK → users, NULL',     'NULL',              'NULL = seeded by system; UUID = owner'),
    ('created_at',   'TIMESTAMP',    'NOT NULL',             'NOW()',             'Waktu pembuatan'),
]
for row in tt_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()
add_para(doc, 'Data Seed Awal (is_system = TRUE):', bold=True)
seed_t = doc.add_table(rows=1, cols=5)
seed_t.style = 'Table Grid'
add_table_header(seed_t, ['code', 'label', 'wallet_type', 'direction', 'Keterangan'])
seed_data = [
    ('owner_topup',       'Top Up Saldo',        'jenius / emoney', 'credit', 'Owner isi saldo'),
    ('owner_withdrawal',  'Penarikan Saldo',      'jenius',          'debit',  'Owner tarik saldo (jarang)'),
    ('atm_withdrawal',    'Tarik ATM',            'jenius',          'debit',  'Driver ambil tunai; auto credit ke Cash'),
    ('fuel_payment',      'Bayar Bensin',         'jenius',          'debit',  'Driver bayar BBM'),
    ('cash_in',           'Uang Cash Masuk',      'cash',            'credit', 'Otomatis dari atm_withdrawal'),
    ('cash_out',          'Pengeluaran Cash',     'cash',            'debit',  'Driver catat pengeluaran tunai'),
    ('emoney_usage',      'Pemakaian eMoney',     'emoney',          'debit',  'Tol, parkir gedung, dll'),
]
for row in seed_data:
    add_table_row(seed_t, list(row), font_size=9)

doc.add_paragraph()

# ── Tabel wallets ────────────────────────────────────────────────────────────
add_heading(doc, '6.5 Tabel: wallets', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
w_cols = [
    ('id',               'UUID',         'PK',                    'gen_random_uuid()', 'Primary key'),
    ('driver_id',        'UUID',         'FK → users, NOT NULL',  '-',                 'Hanya user dengan role=driver'),
    ('wallet_type',      'VARCHAR(20)',  'NOT NULL',              '-',                 "Nilai: 'jenius', 'emoney', 'cash'"),
    ('label',            'VARCHAR(100)', 'NOT NULL',              '-',                 'e.g. "Jenius – Budi"'),
    ('current_balance',  'BIGINT',       'NOT NULL',              '0',                 'Saldo terkini dalam Rupiah (integer)'),
    ('is_active',        'BOOLEAN',      'NOT NULL',              'TRUE',              'FALSE = wallet dinonaktifkan'),
    ('created_at',       'TIMESTAMP',    'NOT NULL',              'NOW()',             'Waktu pembuatan wallet'),
]
for row in w_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_paragraph()

# ── Tabel transactions ───────────────────────────────────────────────────────
add_heading(doc, '6.6 Tabel: transactions', 2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
add_table_header(t, ['Kolom', 'Tipe Data', 'Constraint', 'Default', 'Keterangan'])
tx_cols = [
    ('id',              'UUID',      'PK',                          'gen_random_uuid()', 'Primary key'),
    ('wallet_id',       'UUID',      'FK → wallets, NOT NULL',      '-',                 'Wallet yang terdampak'),
    ('txn_type_id',     'UUID',      'FK → transaction_types, NOT NULL', '-',            'Jenis transaksi'),
    ('created_by',      'UUID',      'FK → users, NOT NULL',        '-',                 'User yang menginput'),
    ('direction',       'VARCHAR(10)','NOT NULL',                   '-',                 "Nilai: 'credit' atau 'debit'"),
    ('amount',          'BIGINT',    'NOT NULL, CHECK > 0',         '-',                 'Nominal selalu positif (Rupiah)'),
    ('balance_before',  'BIGINT',    'NOT NULL',                    '-',                 'Snapshot saldo sebelum transaksi'),
    ('balance_after',   'BIGINT',    'NOT NULL',                    '-',                 'Snapshot saldo setelah transaksi'),
    ('notes',           'TEXT',      'NULL',                        'NULL',              'Catatan bebas'),
    ('txn_date',        'TIMESTAMP', 'NOT NULL',                    '-',                 'Tanggal & jam transaksi (user-input)'),
    ('linked_txn_id',   'UUID',      'FK → transactions, NULL',     'NULL',              'Pasangan linked transaction (ATM↔Cash)'),
    ('created_at',      'TIMESTAMP', 'NOT NULL',                    'NOW()',             'Waktu input ke sistem'),
]
for row in tx_cols:
    add_table_row(t, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — MENU & NAVIGATION
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Struktur Menu & Navigasi', 1)
hr(doc)

add_heading(doc, '7.1 Menu Owner', 2)

menu_owner = doc.add_table(rows=1, cols=3)
menu_owner.style = 'Table Grid'
add_table_header(menu_owner, ['Menu Utama', 'Sub-Menu', 'Fungsi'])
owner_menu_data = [
    ('Dashboard',    '-',                        'Ringkasan saldo semua wallet semua driver + 5 transaksi terakhir'),
    ('Driver',       'Daftar Driver',            'Tabel semua driver beserta status aktif dan ringkasan saldo'),
    ('',             'Tambah Driver',            'Form tambah driver baru + kirim magic link onboarding'),
    ('',             'Edit Driver',              'Ubah nama, status aktif/nonaktif driver'),
    ('Wallet',       'Daftar Wallet',            'Semua wallet dikelompokkan per driver'),
    ('',             'Top Up Jenius',            'Form top up saldo Jenius untuk driver tertentu'),
    ('',             'Top Up eMoney',            'Form top up saldo eMoney untuk driver tertentu'),
    ('',             'Penarikan Jenius',         'Form penarikan saldo Jenius (jarang)'),
    ('Jenis Transaksi', 'Daftar Jenis Transaksi','Tabel semua transaction_types + status is_system'),
    ('',             'Tambah Jenis Transaksi',   'Form tambah jenis transaksi baru'),
    ('',             'Edit Jenis Transaksi',     'Ubah label, sort order, status aktif'),
    ('Laporan',      'Transaksi Jenius',         'Filter per driver & tanggal; tampilkan riwayat transaksi Jenius'),
    ('',             'Laporan Cash',             'Saldo awal, transaksi masuk/keluar, saldo akhir per driver'),
    ('',             'Rekap eMoney',             'Riwayat top up dan pemakaian eMoney per driver'),
    ('Pengaturan',   'Kelola Session Driver',    'Daftar session aktif driver + tombol terminate'),
    ('',             'Logout',                   'Akhiri session owner'),
]
for row_data in owner_menu_data:
    add_table_row(menu_owner, list(row_data), font_size=9)

doc.add_paragraph()

add_heading(doc, '7.2 Menu Driver (Mobile-First)', 2)

menu_driver = doc.add_table(rows=1, cols=3)
menu_driver.style = 'Table Grid'
add_table_header(menu_driver, ['Menu Utama', 'Sub-Menu', 'Fungsi'])
driver_menu_data = [
    ('Dashboard',    '-',                       '3 kartu saldo: Jenius, eMoney, Cash + 5 transaksi terakhir'),
    ('Jenius',       'Catat Transaksi Jenius',  'Form dengan pilihan: Tarik ATM / Bayar Bensin / Lainnya'),
    ('Uang Cash',    'Catat Pengeluaran Cash',  'Form input pengeluaran tunai: nominal, tanggal/jam, catatan'),
    ('Riwayat',      'Transaksi Jenius',        'Daftar riwayat transaksi wallet Jenius milik driver'),
    ('',             'Transaksi Cash',          'Daftar riwayat transaksi wallet Cash milik driver'),
    ('',             'Logout',                  'Akhiri session driver'),
]
for row_data in driver_menu_data:
    add_table_row(menu_driver, list(row_data), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BUSINESS RULES
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '8. Aturan Bisnis', 1)
hr(doc)

br_table = doc.add_table(rows=1, cols=3)
br_table.style = 'Table Grid'
add_table_header(br_table, ['ID', 'Aturan Bisnis', 'Dampak Pelanggaran'])
br_data = [
    ('BR-01', 'Saldo wallet tidak boleh negatif setelah transaksi debit',
     'Transaksi ditolak, tampilkan pesan error "Saldo tidak mencukupi"'),
    ('BR-02', 'Transaksi Tarik ATM harus membuat dua record sekaligus: debit Jenius dan kredit Cash (linked transaction)',
     'Jika salah satu gagal, keduanya di-rollback (atomic)'),
    ('BR-03', 'Nominal transaksi harus lebih besar dari 0 (nol)',
     'Form validasi menolak input dengan nominal ≤ 0'),
    ('BR-04', 'Magic link login hanya berlaku 15 menit dan satu kali pakai',
     'Link kedaluwarsa diarahkan ke halaman "Link tidak valid, minta ulang"'),
    ('BR-05', 'Magic link onboarding berlaku 24 jam dan satu kali pakai',
     'Sama dengan BR-04; owner harus mengirim ulang link'),
    ('BR-06', 'Driver yang dinonaktifkan tidak dapat login meskipun memiliki magic link aktif',
     'Redirect ke halaman "Akun tidak aktif"'),
    ('BR-07', 'Jenis transaksi dengan is_system = TRUE tidak dapat dihapus dan tidak dapat diubah kode serta arahnya',
     'Tombol hapus disembunyikan; form edit mengunci field code dan direction'),
    ('BR-08', 'Jenis transaksi dengan is_system = FALSE tidak dapat dihapus jika sudah digunakan di minimal 1 transaksi',
     'Tampilkan pesan error "Jenis transaksi sudah digunakan"'),
    ('BR-09', 'Setiap wallet memiliki tepat satu driver pemilik',
     'Wallet tidak dapat dipindah ke driver lain setelah dibuat'),
    ('BR-10', 'balance_after harus sama dengan: balance_before + amount (jika credit) atau balance_before – amount (jika debit)',
     'Validasi server-side sebelum INSERT; jika tidak konsisten, transaksi ditolak'),
    ('BR-11', 'Owner tidak dapat menghapus session miliknya sendiri via halaman kelola session',
     'Tombol terminate tidak ditampilkan untuk session milik owner'),
    ('BR-12', 'Driver hanya dapat melihat data miliknya sendiri, tidak bisa melihat data driver lain',
     'Query selalu difilter berdasarkan driver_id = current_user.id'),
]
for row in br_data:
    add_table_row(br_table, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ARSITEKTUR & STACK TEKNOLOGI
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '9. Arsitektur & Stack Teknologi', 1)
hr(doc)

add_heading(doc, '9.1 Stack Teknologi', 2)
tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
add_table_header(tech_table, ['Komponen', 'Teknologi', 'Keterangan'])
tech_data = [
    ('Backend Framework',   'Python Flask',                         'Monolithic MPA; menggunakan Blueprint per modul'),
    ('ORM',                 'Flask-SQLAlchemy + psycopg2-binary',   'Akses database PostgreSQL via abstraksi ORM'),
    ('Database (Prod)',     'PostgreSQL',                           'Database utama untuk production'),
    ('Database (Dev)',      'SQLite',                               'Database lokal untuk development, tidak perlu install'),
    ('Frontend',            'HTML5 + CSS3 + JavaScript Vanilla',    'Tanpa framework CSS; tampilan custom sepenuhnya'),
    ('Partial Update',      'HTMX',                                 'Partial page update tanpa full reload untuk form dan tabel'),
    ('Font',                'Plus Jakarta Sans + JetBrains Mono',   'Via Google Fonts CDN'),
    ('Session',             'Flask Server-Side Session',            'Cookie HTTPONLY + SAMESITE=Lax; session ID di server'),
    ('Email',               'Flask-Mail / SMTP',                    'Pengiriman magic link via email'),
    ('Password/Hash',       'Werkzeug Security',                    'pbkdf2:sha256 dengan 600.000 iterasi (untuk hash token)'),
]
for row in tech_data:
    add_table_row(tech_table, list(row), font_size=9)

doc.add_paragraph()

add_heading(doc, '9.2 Struktur Direktori', 2)
dir_text = (
    "flask_uangkas/\n"
    "├── app.py                    # Application factory (create_app)\n"
    "├── config.py                 # DevelopmentConfig, ProductionConfig\n"
    "├── extensions.py             # Inisialisasi db, mail\n"
    "├── models/\n"
    "│   ├── user.py               # Model User\n"
    "│   ├── session_token.py      # Model UserSession\n"
    "│   ├── magic_token.py        # Model MagicToken\n"
    "│   ├── wallet.py             # Model Wallet\n"
    "│   ├── transaction_type.py   # Model TransactionType\n"
    "│   └── transaction.py        # Model Transaction\n"
    "├── routes/\n"
    "│   ├── auth.py               # Blueprint: login, verify, logout\n"
    "│   ├── owner/\n"
    "│   │   ├── dashboard.py      # Blueprint: owner dashboard\n"
    "│   │   ├── drivers.py        # Blueprint: CRUD driver\n"
    "│   │   ├── wallets.py        # Blueprint: topup, kelola wallet\n"
    "│   │   ├── txn_types.py      # Blueprint: kelola jenis transaksi\n"
    "│   │   ├── reports.py        # Blueprint: laporan & monitoring\n"
    "│   │   └── sessions.py       # Blueprint: kelola session driver\n"
    "│   └── driver/\n"
    "│       ├── dashboard.py      # Blueprint: driver dashboard\n"
    "│       ├── jenius.py         # Blueprint: catat transaksi Jenius\n"
    "│       └── cash.py           # Blueprint: catat pengeluaran cash\n"
    "├── templates/\n"
    "│   ├── base_owner.html       # Layout owner\n"
    "│   ├── base_driver.html      # Layout driver (mobile-first)\n"
    "│   ├── auth/\n"
    "│   ├── owner/\n"
    "│   └── driver/\n"
    "└── static/\n"
    "    ├── css/style.css\n"
    "    └── js/main.js"
)

p = doc.add_paragraph()
run = p.add_run(dir_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — URL ROUTES
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '10. Daftar URL Routes', 1)
hr(doc)

route_table = doc.add_table(rows=1, cols=4)
route_table.style = 'Table Grid'
add_table_header(route_table, ['Method', 'URL', 'Akses', 'Fungsi'])
route_data = [
    # Auth
    ('GET/POST', '/login',                          'Public',  'Form email → kirim magic link'),
    ('GET',      '/auth/verify/<token>',            'Public',  'Verifikasi magic link, buat session'),
    ('POST',     '/auth/logout',                    'Login',   'Hapus session aktif'),
    # Owner
    ('GET',      '/owner/',                         'Owner',   'Dashboard ringkasan semua wallet'),
    ('GET',      '/owner/drivers',                  'Owner',   'Daftar driver'),
    ('GET/POST', '/owner/drivers/add',              'Owner',   'Form tambah driver'),
    ('GET/POST', '/owner/drivers/<id>/edit',        'Owner',   'Edit nama & status driver'),
    ('POST',     '/owner/drivers/<id>/resend',      'Owner',   'Kirim ulang magic link onboarding'),
    ('GET',      '/owner/wallets',                  'Owner',   'Daftar semua wallet'),
    ('POST',     '/owner/wallets/<id>/topup',       'Owner',   'Top up saldo Jenius/eMoney'),
    ('POST',     '/owner/wallets/<id>/withdraw',    'Owner',   'Penarikan saldo Jenius'),
    ('GET',      '/owner/txn-types',                'Owner',   'Daftar jenis transaksi'),
    ('GET/POST', '/owner/txn-types/add',            'Owner',   'Tambah jenis transaksi baru'),
    ('GET/POST', '/owner/txn-types/<id>/edit',      'Owner',   'Edit jenis transaksi'),
    ('POST',     '/owner/txn-types/<id>/delete',    'Owner',   'Hapus jenis transaksi (non-system, belum dipakai)'),
    ('GET',      '/owner/reports/jenius',           'Owner',   'Laporan transaksi Jenius'),
    ('GET',      '/owner/reports/cash',             'Owner',   'Laporan uang cash'),
    ('GET',      '/owner/reports/emoney',           'Owner',   'Rekap eMoney'),
    ('GET',      '/owner/sessions',                 'Owner',   'Daftar session aktif driver'),
    ('POST',     '/owner/sessions/<id>/terminate',  'Owner',   'Terminate session driver'),
    # Driver
    ('GET',      '/driver/',                        'Driver',  'Dashboard 3 kartu saldo'),
    ('GET/POST', '/driver/jenius/record',           'Driver',  'Form catat transaksi Jenius'),
    ('GET/POST', '/driver/cash/record',             'Driver',  'Form catat pengeluaran cash'),
    ('GET',      '/driver/history/jenius',          'Driver',  'Riwayat transaksi Jenius'),
    ('GET',      '/driver/history/cash',            'Driver',  'Riwayat transaksi Cash'),
]
for row in route_data:
    add_table_row(route_table, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — USE CASE FLOWS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '11. Alur Use Case Utama', 1)
hr(doc)

add_heading(doc, 'UC-01: Login dengan Magic Link', 2)
uc1_steps = [
    'Pengguna membuka aplikasi dan diarahkan ke halaman /login.',
    'Pengguna memasukkan alamat email dan menekan tombol "Kirim Magic Link".',
    'Sistem memvalidasi: email terdaftar? akun aktif?',
    '[Gagal] Jika email tidak terdaftar atau akun nonaktif → tampilkan pesan error yang generik (tidak spesifik untuk keamanan).',
    '[Sukses] Sistem membuat token unik, meng-hash-nya, menyimpan ke tabel magic_tokens, lalu mengirim email dengan magic link.',
    'Pengguna mengklik magic link di email.',
    'Sistem mencari token (setelah di-hash), memvalidasi: belum kedaluwarsa? belum pernah dipakai?',
    '[Gagal] Token tidak valid → redirect ke /login dengan pesan "Link tidak valid atau sudah kedaluwarsa".',
    '[Sukses] Sistem menandai token sebagai terpakai (used_at = NOW()), membuat record di user_sessions, menyimpan session_token di cookie, lalu redirect ke dashboard sesuai peran.',
]
for i, step in enumerate(uc1_steps, 1):
    add_numbered(doc, step)

doc.add_paragraph()

add_heading(doc, 'UC-02: Driver Mencatat Tarik ATM', 2)
uc2_steps = [
    'Driver login dan membuka menu Jenius → Catat Transaksi Jenius.',
    'Driver memilih jenis transaksi "Tarik ATM".',
    'Driver mengisi nominal (e.g. Rp 300.000), tanggal & jam transaksi, dan catatan opsional.',
    'Driver menekan tombol "Simpan Transaksi".',
    'Sistem memvalidasi: saldo Jenius driver ≥ nominal? nominal > 0?',
    '[Gagal Validasi] Tampilkan pesan error, form tidak disimpan.',
    '[Sukses] Sistem memulai database transaction.',
    'INSERT ke tabel transactions: debit Jenius, amount=300.000, balance_before=saldo_lama, balance_after=saldo_lama-300.000.',
    'UPDATE wallets SET current_balance = current_balance - 300.000 WHERE id = jenius_wallet_id.',
    'INSERT ke tabel transactions: credit Cash, amount=300.000, linked_txn_id=id_dari_langkah_8.',
    'UPDATE wallets SET current_balance = current_balance + 300.000 WHERE id = cash_wallet_id.',
    'COMMIT database transaction.',
    '[Jika error di langkah 8–11] ROLLBACK; tampilkan pesan error generik.',
    'Redirect ke halaman riwayat atau dashboard dengan pesan sukses.',
]
for i, step in enumerate(uc2_steps, 1):
    add_numbered(doc, step)

doc.add_paragraph()

add_heading(doc, 'UC-03: Owner Terminate Session Driver', 2)
uc3_steps = [
    'Owner membuka menu Pengaturan → Kelola Session Driver.',
    'Halaman menampilkan daftar session aktif semua driver (is_active = TRUE).',
    'Owner mengklik tombol "Terminate" pada session driver tertentu.',
    'Sistem mengupdate: user_sessions SET is_active=FALSE, terminated_at=NOW(), terminated_by=owner_id WHERE id=session_id.',
    'Pada request berikutnya dari driver tersebut, sistem menemukan session_token di cookie, mencari di tabel user_sessions, mendapati is_active=FALSE, lalu menghapus cookie dan redirect ke /login.',
]
for i, step in enumerate(uc3_steps, 1):
    add_numbered(doc, step)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — CONSTRAINTS & ASSUMPTIONS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '12. Batasan & Asumsi', 1)
hr(doc)

add_heading(doc, '12.1 Batasan (Constraints)', 2)
constraints = [
    'Aplikasi tidak berintegrasi dengan API bank Jenius atau sistem eMoney; semua data diinput manual.',
    'Tidak ada fitur ekspor ke Excel atau PDF pada versi 1.0.',
    'Tidak ada notifikasi real-time (push notification atau SMS); hanya email untuk magic link.',
    'Jumlah driver tidak dibatasi secara teknis, namun desain awal ditujukan untuk 2 driver.',
    'Aplikasi tidak mengelola penggajian atau slip gaji driver.',
    'Tidak ada backup otomatis; backup database menjadi tanggung jawab infrastruktur hosting.',
]
for c in constraints:
    add_bullet(doc, c)

doc.add_paragraph()

add_heading(doc, '12.2 Asumsi (Assumptions)', 2)
assumptions = [
    'Owner memiliki akses email setyanaputra@yahoo.com dan dapat menerima email dari sistem.',
    'Setiap driver memiliki smartphone dengan browser modern dan koneksi internet.',
    'SMTP server untuk pengiriman email tersedia dan dikonfigurasi (e.g. Gmail SMTP atau layanan transaksional).',
    'Owner bertanggung jawab atas keakuratan data top up yang diinput (tidak ada rekonsiliasi otomatis dengan bank).',
    'Saldo awal semua wallet saat pertama kali driver ditambahkan adalah Rp 0.',
    'Nominal transaksi menggunakan satuan Rupiah penuh (tanpa sen/desimal).',
    'Satu akun email hanya dapat menjadi satu pengguna (satu peran: owner atau driver).',
]
for a in assumptions:
    add_bullet(doc, a)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — ACCEPTANCE CRITERIA
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '13. Kriteria Penerimaan', 1)
hr(doc)

ac_table = doc.add_table(rows=1, cols=3)
ac_table.style = 'Table Grid'
add_table_header(ac_table, ['ID', 'Skenario', 'Hasil yang Diharapkan'])
ac_data = [
    ('AC-01', 'Owner memasukkan email valid dan mengklik magic link',
     'Owner masuk ke dashboard owner; session aktif tanpa batas waktu'),
    ('AC-02', 'Driver memasukkan email valid dan mengklik magic link',
     'Driver masuk ke dashboard driver mobile; saldo tiga wallet ditampilkan'),
    ('AC-03', 'Driver mencatat tarik ATM Rp 200.000 dengan saldo Jenius Rp 500.000',
     'Saldo Jenius berkurang menjadi Rp 300.000; saldo Cash bertambah Rp 200.000; dua record transaksi terhubung'),
    ('AC-04', 'Driver mencatat tarik ATM Rp 600.000 dengan saldo Jenius Rp 500.000',
     'Transaksi ditolak; pesan error "Saldo tidak mencukupi" ditampilkan'),
    ('AC-05', 'Owner melakukan top up Jenius driver Rp 1.000.000',
     'Saldo Jenius driver bertambah Rp 1.000.000; transaksi tercatat dengan tipe owner_topup'),
    ('AC-06', 'Owner menterminasi session driver aktif',
     'Session driver di-set is_active=FALSE; driver diredirect ke /login pada request berikutnya'),
    ('AC-07', 'Owner menambah jenis transaksi baru "Bayar Cuci Mobil" untuk wallet Jenius',
     'Jenis transaksi muncul di dropdown driver saat catat transaksi Jenius'),
    ('AC-08', 'Owner mencoba menghapus jenis transaksi dengan is_system=TRUE',
     'Tombol hapus tidak tersedia; jika akses langsung via URL dikembalikan error 403'),
    ('AC-09', 'Magic link yang sama digunakan dua kali',
     'Klik pertama: berhasil login. Klik kedua: "Link tidak valid atau sudah digunakan"'),
    ('AC-10', 'Magic link digunakan setelah 15 menit',
     'Diarahkan ke halaman "Link kedaluwarsa, silakan minta link baru"'),
    ('AC-11', 'Driver yang dinonaktifkan mencoba login',
     'Magic link berhasil dikirim (sistem tidak membocorkan status), namun saat verifikasi ditampilkan pesan "Akun tidak aktif"'),
    ('AC-12', 'Driver mencoba mengakses URL endpoint owner',
     'Redirect ke 403 Forbidden atau dashboard driver'),
]
for row in ac_data:
    add_table_row(ac_table, list(row), font_size=9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — SECURITY REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '14. Kebutuhan Keamanan', 1)
hr(doc)

sec_items = [
    ('SEC-01: Proteksi CSRF',
     'Semua form yang menggunakan metode POST harus menyertakan CSRF token yang divalidasi oleh server sebelum proses data.'),
    ('SEC-02: Cookie Keamanan',
     'Session cookie dikonfigurasi dengan flag HTTPONLY (tidak dapat diakses JavaScript), SAMESITE=Lax (proteksi CSRF lintas domain), dan SECURE=True pada environment production (hanya dikirim via HTTPS).'),
    ('SEC-03: Hashing Token Magic Link',
     'Token magic link yang disimpan di database adalah hasil hash SHA-256; token plaintext hanya dikirim via email dan tidak pernah disimpan.'),
    ('SEC-04: Role-Based Access Control',
     'Setiap endpoint diproteksi dengan decorator yang memvalidasi peran pengguna. Endpoint /owner/* hanya dapat diakses oleh role=owner; endpoint /driver/* hanya oleh role=driver.'),
    ('SEC-05: Validasi Input Server-Side',
     'Semua input pengguna divalidasi di sisi server (bukan hanya di browser) untuk mencegah bypass via request langsung. Validasi mencakup: tipe data, range nilai, panjang string.'),
    ('SEC-06: Proteksi SQL Injection',
     'Seluruh query database menggunakan parameterized query melalui SQLAlchemy ORM; tidak ada string concatenation untuk query.'),
    ('SEC-07: Proteksi XSS',
     'Semua output di template HTML menggunakan auto-escaping Jinja2; tidak ada penggunaan |safe kecuali untuk konten yang benar-benar aman.'),
    ('SEC-08: Rate Limiting Magic Link',
     'Pembatasan maksimal 3 permintaan magic link per email per 5 menit untuk mencegah spam email.'),
    ('SEC-09: Audit Trail Transaksi',
     'Setiap transaksi keuangan menyimpan: user yang menginput, waktu input, saldo sebelum, dan saldo sesudah untuk keperluan audit.'),
    ('SEC-10: Isolasi Data Driver',
     'Query data transaksi dan wallet selalu difilter berdasarkan user yang sedang login; driver tidak dapat mengakses data driver lain meskipun mengetahui ID-nya.'),
]
for title, desc in sec_items:
    add_para(doc, title, bold=True)
    add_para(doc, desc)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — DIAGRAM ERD DESKRIPTIF
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '15. Relasi Antar Entitas (ERD Deskriptif)', 1)
hr(doc)

erd_table = doc.add_table(rows=1, cols=3)
erd_table.style = 'Table Grid'
add_table_header(erd_table, ['Entitas Asal', 'Kardinalitas', 'Entitas Tujuan'])
erd_data = [
    ('users (role=driver)',    '1 ──── N',  'wallets'),
    ('users',                 '1 ──── N',  'magic_tokens'),
    ('users',                 '1 ──── N',  'user_sessions'),
    ('users',                 '1 ──── N',  'transactions (created_by)'),
    ('users (role=owner)',    '1 ──── N',  'transaction_types (created_by)'),
    ('wallets',               '1 ──── N',  'transactions'),
    ('transaction_types',     '1 ──── N',  'transactions (txn_type_id)'),
    ('transactions',          '1 ──── 0/1','transactions (linked_txn_id, self-referencing)'),
    ('users (owner)',         '1 ──── 0/N','user_sessions (terminated_by)'),
]
for row in erd_data:
    add_table_row(erd_table, list(row), align_center_cols=[1], font_size=9)

doc.add_paragraph()
add_para(doc, 'Catatan Kardinalitas:', bold=True)
add_bullet(doc, '1 driver memiliki tepat 3 wallet (Jenius, eMoney, Cash) yang dibuat otomatis saat driver ditambahkan.')
add_bullet(doc, 'Setiap transaksi merujuk ke tepat 1 wallet dan 1 transaction_type.')
add_bullet(doc, 'Transaksi ATM Withdrawal memiliki linked_txn_id yang merujuk ke transaksi Cash In yang merupakan pasangannya.')
add_bullet(doc, 'Setiap magic_token memiliki masa berlaku dan hanya dapat digunakan satu kali.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 16 — SIGN-OFF
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, '16. Persetujuan Dokumen', 1)
hr(doc)

add_para(doc, 'Dokumen ini dianggap sah dan disetujui setelah ditandatangani oleh pihak-pihak berikut:')
doc.add_paragraph()

sign_table = doc.add_table(rows=4, cols=3)
sign_table.style = 'Table Grid'
add_table_header(sign_table, ['Peran', 'Nama', 'Tanda Tangan & Tanggal'])
sign_data = [
    ('Owner / Pemohon',     'Setyana Putra',    ' '),
    ('Lead Developer',      ' ',                ' '),
    ('Quality Assurance',   ' ',                ' '),
]
for i, row_data in enumerate(sign_data):
    row = sign_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.text.strip() else Pt(14)
    # Beri ruang untuk tanda tangan
    for cell in row.cells:
        cell.height = Cm(2.0)

# ─── Save ───────────────────────────────────────────────────────────────────

output_path = r'c:\Python312\flask_uangkas\BRD_UangKas_v1.0.docx'
doc.save(output_path)
print(f"BRD berhasil dibuat: {output_path}")
